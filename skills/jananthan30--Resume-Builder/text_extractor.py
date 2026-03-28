"""
Centralized text extraction for resume and JD files.

Handles:
  - DOCX: paragraphs + tables + text boxes (shapes)
  - PDF (digital): pdfplumber with layout-aware ordering
  - PDF (scanned): OCR via pytesseract (if installed) or Claude Vision API
  - MD / TXT: plain UTF-8 read

Usage:
    from text_extractor import extract_text
    text = extract_text("resume.pdf")
    text = extract_text("resume.docx")
"""

import base64
import io
import json
import os
import re
import urllib.request
from pathlib import Path
from typing import Optional


# ---------------------------------------------------------------------------
# DOCX Extraction
# ---------------------------------------------------------------------------

def _extract_docx(file_path: str) -> str:
    """Extract text from DOCX including tables and text boxes."""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(file_path)
    parts: list[str] = []

    # Paragraphs (main body)
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables — often used for skills grids, two-column layouts
    for table in doc.tables:
        for row in table.rows:
            row_parts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_parts:
                parts.append("  ".join(row_parts))

    # Text boxes / shapes (wps:txbx elements)
    try:
        body = doc.element.body
        for txbx in body.iter(qn("wps:txbx")):
            for para_el in txbx.iter(qn("w:p")):
                text = "".join(
                    r.text for r in para_el.iter(qn("w:r"))
                    if r.text
                )
                if text.strip():
                    parts.append(text.strip())
    except Exception:
        pass  # text boxes are optional; skip if namespace lookup fails

    # Headers and footers
    try:
        for section in doc.sections:
            for hf in [section.header, section.footer]:
                if hf and not hf.is_linked_to_previous:
                    for para in hf.paragraphs:
                        if para.text.strip():
                            parts.append(para.text)
    except Exception:
        pass

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# PDF Extraction — digital (text-based)
# ---------------------------------------------------------------------------

def _extract_pdf_digital(file_path: str) -> str:
    """Extract text from a text-based (digital) PDF using pdfplumber."""
    try:
        import pdfplumber
        pages: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                pages.append(text)
        return "\n".join(pages)
    except ImportError:
        pass

    # Fallback: PyMuPDF
    try:
        import fitz
        doc = fitz.open(file_path)
        pages = [page.get_text() for page in doc]
        doc.close()
        return "\n".join(pages)
    except ImportError:
        pass

    raise ImportError(
        "Install pdfplumber for PDF support: pip install pdfplumber"
    )


def _is_scanned(text: str, page_count: int) -> bool:
    """Return True if the extracted text looks like a scanned (image-only) PDF."""
    if page_count == 0:
        return False
    avg_chars = len(text.strip()) / page_count
    return avg_chars < 80  # fewer than 80 chars/page → likely scanned


# ---------------------------------------------------------------------------
# PDF Extraction — scanned (OCR)
# ---------------------------------------------------------------------------

def _ocr_via_claude(file_path: str) -> str:
    """
    OCR a scanned PDF using Claude's vision API.
    Requires ANTHROPIC_API_KEY env var and PyMuPDF (fitz) for rendering.
    """
    api_key = os.getenv("ANTHROPIC_API_KEY", "")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY not set — cannot OCR scanned PDF")

    try:
        import fitz
    except ImportError:
        raise ImportError("Install PyMuPDF for scanned PDF support: pip install pymupdf")

    model = os.getenv("ANTHROPIC_MODEL", "claude-haiku-4-5-20251001")
    doc = fitz.open(file_path)
    extracted_pages: list[str] = []

    for page_num, page in enumerate(doc):
        # Render at 150 DPI (good enough for OCR, keeps payload small)
        mat = fitz.Matrix(150 / 72, 150 / 72)
        pix = page.get_pixmap(matrix=mat)
        img_bytes = pix.tobytes("png")
        img_b64 = base64.standard_b64encode(img_bytes).decode("utf-8")

        payload = json.dumps({
            "model": model,
            "max_tokens": 2000,
            "messages": [{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": "image/png",
                            "data": img_b64,
                        },
                    },
                    {
                        "type": "text",
                        "text": (
                            "This is page " + str(page_num + 1) + " of a resume. "
                            "Extract ALL text exactly as it appears. "
                            "Preserve line breaks and structure. "
                            "Return only the extracted text, no commentary."
                        ),
                    },
                ],
            }],
        }).encode("utf-8")

        req = urllib.request.Request(
            "https://api.anthropic.com/v1/messages",
            data=payload,
            headers={
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json",
            },
            method="POST",
        )
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        extracted_pages.append(data["content"][0]["text"])

    doc.close()
    return "\n\n".join(extracted_pages)


def _ocr_via_tesseract(file_path: str) -> str:
    """OCR a scanned PDF using pytesseract + pdf2image (local, no API key needed)."""
    try:
        import pytesseract
        from pdf2image import convert_from_path
    except ImportError:
        raise ImportError(
            "Install tesseract OCR support: pip install pytesseract pdf2image\n"
            "Also install Tesseract binary: https://github.com/UB-Mannheim/tesseract/wiki"
        )

    images = convert_from_path(file_path, dpi=150)
    pages = [pytesseract.image_to_string(img) for img in images]
    return "\n\n".join(pages)


def _extract_pdf(file_path: str) -> str:
    """
    Extract text from a PDF. Automatically detects scanned PDFs and
    applies OCR if needed.

    OCR priority:
      1. Claude Vision API (if ANTHROPIC_API_KEY is set) — no extra install
      2. pytesseract + pdf2image (if installed locally)
      3. Return partial text with a warning comment
    """
    # Get page count for scanned detection
    page_count = 1
    try:
        import fitz
        doc = fitz.open(file_path)
        page_count = len(doc)
        doc.close()
    except ImportError:
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
        except Exception:
            pass

    text = _extract_pdf_digital(file_path)

    if not _is_scanned(text, page_count):
        return text

    # Scanned PDF — try OCR
    # Option 1: Claude Vision API
    if os.getenv("ANTHROPIC_API_KEY"):
        try:
            return _ocr_via_claude(file_path)
        except Exception:
            pass

    # Option 2: pytesseract
    try:
        return _ocr_via_tesseract(file_path)
    except ImportError:
        pass

    # Option 3: Return what we have with a warning
    return (
        "[WARNING: This appears to be a scanned PDF. Text extraction is limited.]\n"
        "[To enable OCR: set ANTHROPIC_API_KEY, or install pytesseract + pdf2image.]\n\n"
        + text
    )


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_text(file_path: str) -> str:
    """
    Extract plain text from a resume or JD file.

    Supported formats: .pdf, .docx, .md, .txt
    Handles: digital PDFs, scanned PDFs (OCR), DOCX tables, DOCX text boxes.
    """
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    elif ext in (".md", ".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    else:
        raise ValueError(
            f"Unsupported file format: {ext}. "
            "Supported: .pdf, .docx, .md, .txt"
        )


# Keep backward-compatible alias used in existing code
extract_text_from_file = extract_text
