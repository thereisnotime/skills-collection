from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from docx_generator import (
    _assert_docx_list_item_parity,
    create_ats_resume,
    parse_resume_markdown,
)


def _generate_resume(path: Path, summary: str) -> Document:
    create_ats_resume(
        output_path=str(path),
        name="Jane Doe, MD",
        contact_info={"email": "jane@example.com"},
        summary=summary,
        core_competencies=["Medical Safety Review"],
        experience=[],
        education=[],
        certifications=[],
    )
    return Document(path)


def _has_bottom_border(paragraph) -> bool:
    properties = paragraph._p.pPr
    if properties is None:
        return False
    borders = properties.find(qn("w:pBdr"))
    return borders is not None and borders.find(qn("w:bottom")) is not None


@pytest.mark.parametrize("summary", ["", "   \t  "])
def test_blank_summary_omits_heading_paragraph_and_extra_separator(
    tmp_path: Path, summary: str
) -> None:
    document = _generate_resume(tmp_path / "blank-summary.docx", summary)
    paragraphs = document.paragraphs
    texts = [paragraph.text for paragraph in paragraphs]

    assert "PROFESSIONAL SUMMARY" not in texts

    core_index = texts.index("CORE COMPETENCIES")
    assert _has_bottom_border(paragraphs[core_index - 1])
    assert paragraphs[core_index - 2].text == "jane@example.com"


def test_nonblank_summary_keeps_existing_section_and_separator(tmp_path: Path) -> None:
    summary = "Reviewed clinical safety data."
    document = _generate_resume(tmp_path / "nonblank-summary.docx", summary)
    paragraphs = document.paragraphs
    texts = [paragraph.text for paragraph in paragraphs]

    summary_index = texts.index("PROFESSIONAL SUMMARY")
    assert texts[summary_index + 1] == summary
    assert _has_bottom_border(paragraphs[summary_index + 2])
    assert texts[summary_index + 3] == "CORE COMPETENCIES"


def test_heading_only_markdown_preserves_authorized_list_items(tmp_path: Path) -> None:
    markdown = """JANE DOE, MD
New York, NY  |  jane@example.com  |  linkedin.com/in/jane-doe

CORE COMPETENCIES

• Medical Safety Review

PROFESSIONAL EXPERIENCE

MEDICAL REVIEWER | Example Biotech | New York, NY
2020 - Present

• Reviewed safety cases.
"""
    data = parse_resume_markdown(markdown)

    assert data["core_competencies"] == ["Medical Safety Review"]
    assert data["experience"][0]["bullets"] == ["Reviewed safety cases."]
    assert data["contact_info"]["linkedin"] == "linkedin.com/in/jane-doe"

    output_path = tmp_path / "heading-only.docx"
    create_ats_resume(
        output_path=str(output_path),
        name=data["name"],
        contact_info=data["contact_info"],
        summary=data["summary"],
        core_competencies=data["core_competencies"],
        experience=data["experience"],
        education=data["education"],
        certifications=data["certifications"],
    )

    _assert_docx_list_item_parity(markdown, str(output_path))
    rendered_text = [paragraph.text for paragraph in Document(output_path).paragraphs]
    assert "linkedin.com/in/jane-doe" in rendered_text
    assert "Medical Safety Review" in rendered_text
    assert "Reviewed safety cases." in rendered_text


def test_pipe_delimited_competencies_reconstruct_native_bullet_rows(
    tmp_path: Path,
) -> None:
    markdown = """JANE DOE, MD
New York, NY  |  jane@example.com

CORE COMPETENCIES

• Accuracy | Credibility | Reference Selection
• Safety Surveillance | Risk Management | Medical Review
• ICH-GCP | Cross-Functional Teams
"""
    data = parse_resume_markdown(markdown)

    assert data["core_competencies"] == [
        "Accuracy",
        "Credibility",
        "Reference Selection",
        "Safety Surveillance",
        "Risk Management",
        "Medical Review",
        "ICH-GCP",
        "Cross-Functional Teams",
    ]

    output_path = tmp_path / "competency-rows.docx"
    create_ats_resume(
        output_path=str(output_path),
        name=data["name"],
        contact_info=data["contact_info"],
        summary=data["summary"],
        core_competencies=data["core_competencies"],
        experience=[],
        education=[],
        certifications=[],
    )

    _assert_docx_list_item_parity(markdown, str(output_path))
    document = Document(output_path)
    competency_rows = [
        " ".join(paragraph.text.split())
        for paragraph in document.paragraphs
        if paragraph.style.name == "List Bullet"
    ]
    assert competency_rows == [
        "Accuracy | Credibility | Reference Selection",
        "Safety Surveillance | Risk Management | Medical Review",
        "ICH-GCP | Cross-Functional Teams",
    ]


def test_variable_competency_rows_and_certifications_training_preserve_items(
    tmp_path: Path,
) -> None:
    markdown = """JANE DOE, MD
New York, NY  |  jane@example.com

CORE COMPETENCIES

• Accuracy | Credibility | Reference Selection | Safety Surveillance

CERTIFICATIONS & TRAINING

• ECFMG Certified; USMLE Steps 1, 2 CK, and 3 passed
• ICH-GCP trained; Clinical Trials Operations - Johns Hopkins University
"""
    data = parse_resume_markdown(markdown)

    assert data["core_competencies"] == [
        "Accuracy",
        "Credibility",
        "Reference Selection",
        "Safety Surveillance",
    ]
    assert data["certifications"] == [
        "ECFMG Certified; USMLE Steps 1, 2 CK, and 3 passed",
        "ICH-GCP trained; Clinical Trials Operations - Johns Hopkins University",
    ]

    output_path = tmp_path / "variable-competencies-and-training.docx"
    create_ats_resume(
        output_path=str(output_path),
        name=data["name"],
        contact_info=data["contact_info"],
        summary=data["summary"],
        core_competencies=data["core_competencies"],
        experience=[],
        education=[],
        certifications=data["certifications"],
    )

    _assert_docx_list_item_parity(markdown, str(output_path))
    rendered_text = [paragraph.text for paragraph in Document(output_path).paragraphs]
    assert "Accuracy  |  Credibility  |  Reference Selection" in rendered_text
    assert "Safety Surveillance" in rendered_text
    assert data["certifications"][0] in rendered_text
    assert data["certifications"][1] in rendered_text


@pytest.mark.parametrize("location", ["New York, NY", ""])
def test_experience_header_stays_with_first_bullet(
    tmp_path: Path, location: str
) -> None:
    output_path = tmp_path / f"experience-keep-{bool(location)}.docx"
    create_ats_resume(
        output_path=str(output_path),
        name="Jane Doe, MD",
        contact_info={"email": "jane@example.com"},
        summary="",
        core_competencies=["Medical Safety Review"],
        experience=[
            {
                "title": "MEDICAL REVIEWER",
                "company": "Example Biotech",
                "location": location,
                "dates": "2020 - Present",
                "bullets": ["Reviewed safety cases."],
            }
        ],
        education=[],
        certifications=[],
    )
    document = Document(output_path)
    title = next(p for p in document.paragraphs if p.text.startswith("MEDICAL REVIEWER"))
    company = next(p for p in document.paragraphs if p.text == "Example Biotech")
    first_bullet = next(p for p in document.paragraphs if p.text == "Reviewed safety cases.")

    assert title.paragraph_format.keep_with_next is True
    assert company.paragraph_format.keep_with_next is True
    if location:
        location_paragraph = next(p for p in document.paragraphs if p.text == location)
        assert location_paragraph.paragraph_format.keep_with_next is True
    assert first_bullet.paragraph_format.keep_with_next is not True


def test_experience_and_education_date_tabs_use_full_content_width(
    tmp_path: Path,
) -> None:
    output_path = tmp_path / "right-tabs.docx"
    create_ats_resume(
        output_path=str(output_path),
        name="Jane Doe, MD",
        contact_info={"email": "jane@example.com"},
        summary="",
        core_competencies=["Medical Safety Review"],
        experience=[
            {
                "title": "MEDICAL REVIEWER",
                "company": "Example Biotech",
                "location": "New York, NY",
                "dates": "2020 - Present",
                "bullets": ["Reviewed safety cases."],
            }
        ],
        education=[
            {
                "degree": "Doctor of Medicine",
                "school": "Example University",
                "location": "New York, NY",
                "dates": "2010 - 2014",
            }
        ],
        certifications=[],
    )
    document = Document(output_path)
    dated_paragraphs = [
        next(p for p in document.paragraphs if p.text.startswith("MEDICAL REVIEWER")),
        next(p for p in document.paragraphs if p.text.startswith("Doctor of Medicine")),
    ]

    for paragraph in dated_paragraphs:
        tabs = paragraph._p.pPr.find(qn("w:tabs"))
        assert tabs is not None
        tab = tabs.find(qn("w:tab"))
        assert tab is not None
        assert tab.get(qn("w:val")) == "right"
        assert tab.get(qn("w:pos")) == "10224"
