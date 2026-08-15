import json
import shutil
import subprocess
import sys
from pathlib import Path

import pytest
from docx import Document

import docx_generator
import tracker_utils
from candidate_fit_preflight import (
    assess_candidate_fit as _real_assess_candidate_fit,
)
from evidence_engine.testing import DeterministicJudge as _DeterministicJudge


def assess_candidate_fit(*args, **kwargs):
    """Run the fit gate offline with the package's deterministic judge.

    The gate calls a hosted model in production and fails closed without one.
    Tests inject the rule-based judge so pipeline wiring stays testable
    without an API key; production never falls back this way.
    """
    kwargs.setdefault("llm", _DeterministicJudge())
    return _real_assess_candidate_fit(*args, **kwargs)

from final_receipt_verifier import (
    FinalReceiptVerificationError,
    canonical_digest,
    verify_final_receipt,
)
from legacy_rewrite_guard import (
    NATIVE_RESUME_TEAM_REQUIRED,
    NativeResumeTeamRequiredError,
)

ROOT = Path(__file__).resolve().parents[1]

MATCHING_JOB_DESCRIPTION = """Signal Keeper
About the job
Review paper signal records and write findings for safety teams. Manage data and document signal detection work using clear medical writing across paper safety operations.
Qualifications
Experience with Signal Detection, Risk Management, Medical Writing, and Data Management. Review records, assess signals, and write findings for paper safety teams and internal stakeholders.
"""

KNOCKOUT_JOB_DESCRIPTION = """Director, Drug Safety Physician
About the job
Lead safety surveillance for oncology products and provide medical oversight for aggregate reporting and signal management across a global biotech portfolio.
Qualifications
20+ years' experience in Drug Safety/Pharmacovigilance in a biotech company.
Medical Degree (MD) with medical practice experience.
Strong knowledge of ICH and CIOMS guidelines.
Experience reviewing ICSRs and preparing DSUR and PBRER reports.
"""


def _canonical_json(value):
    return json.dumps(
        value,
        sort_keys=True,
        separators=(",", ":"),
        ensure_ascii=True,
    )


def _authorized_artifacts(
    tmp_path: Path,
    resume_text: str | None = None,
    *,
    master_path: Path | None = None,
    job_description: str = MATCHING_JOB_DESCRIPTION,
):
    if resume_text is None:
        resume_text = _canonical_resume()
    master = master_path or (tmp_path / "source-master.md")
    if not master.exists():
        master.write_text(resume_text, encoding="utf-8")
    config = tmp_path / "config.json"
    config.write_text(
        json.dumps({"master_resume_path": str(master)}), encoding="utf-8"
    )
    (tmp_path / "job_description.txt").write_text(
        job_description, encoding="utf-8"
    )
    resume = tmp_path / "resume.md"
    resume.write_text(resume_text, encoding="utf-8")
    draft_digest = canonical_digest(resume.read_bytes())
    votes = [
        {
            "schema_version": "resume-team-vote/v1",
            "name": name,
            "invocation_id": f"vote-{index}",
            "passed": True,
            "draft_digest": draft_digest,
            "codes": [],
        }
        for index, name in enumerate(
            ("evidence", "human_voice", "canonical_integrity"), start=1
        )
    ]
    report = {
        "schema_version": "resume-team-authorization/v1",
        "draft_digest": draft_digest,
        "passed": True,
        "codes": [],
        "findings": [],
        "votes": votes,
    }
    run_id = "run-guard"
    case_id = "case-guard"
    candidate_fit_report = assess_candidate_fit(
        master.read_text(encoding="utf-8"),
        job_description,
        run_id=run_id,
        case_id=case_id,
        as_of_date="2026-07-19",
    )
    assert candidate_fit_report["passed"] is True
    receipt = {
        "schema_version": "resume-team-final-receipt/v2",
        "run_id": run_id,
        "case_id": case_id,
        "draft_digest": draft_digest,
        "source_digest": canonical_digest(master.read_text(encoding="utf-8")),
        "job_description_digest": canonical_digest(job_description),
        "candidate_fit_report": candidate_fit_report,
        "candidate_fit_report_digest": canonical_digest(candidate_fit_report),
        "candidate_fit_judge_report": None,
        "candidate_fit_judge_report_digest": "",
        "researcher_agent_id": "codex:researcher-guard",
        "researcher_artifact_digest": "b" * 64,
        "auditor_attestation": {
            "agent_id": "codex:auditor-guard",
            "artifact_digest": "a" * 64,
            "verdict": "PASS",
            "draft_digest": draft_digest,
        },
        "authorization_digest": canonical_digest(report),
        "authorization_report": report,
        "vote_invocation_ids": [vote["invocation_id"] for vote in votes],
        "publication_id": "publication:guard",
        "verified_target_digest": draft_digest,
    }
    pair_digest = canonical_digest({"run_id": run_id, "case_id": case_id})
    receipt_path = tmp_path / f"resume-team-receipt.{pair_digest}.json"
    receipt_path.write_text(_canonical_json(receipt), encoding="utf-8")
    return resume, receipt_path, receipt, canonical_digest(receipt)


def _canonical_resume(*, title="Signal Keeper"):
    separator = "_" * 79
    return f"""Synthetic Person
Neutral, ZZ | neutral@example.invalid

{separator}
PROFESSIONAL EXPERIENCE

{title} | Paper Harbor Co.
2020 – Present

• Reviewed signal records and wrote plain findings.

{separator}
CORE COMPETENCIES

Signal Detection | Risk Management | Medical Writing | Data Management

{separator}
EDUCATION

Diploma in Paper Systems
Fixture School, Neutral, ZZ | 2018

{separator}
CERTIFICATIONS & LICENSURE

• Paper Safety Permit | PSP-1

{separator}
PUBLICATIONS

1. Reed A. "Paper Signals." Fixture Journal. 2024;1:1-2.

Full publication list: https://example.invalid/paper

{separator}
PROFESSIONAL MEMBERSHIPS

• Society of Paper Fixtures
"""


def test_valid_receipt_verifies_public_schema_and_cross_digests(tmp_path):
    resume, receipt_path, receipt, expected_digest = _authorized_artifacts(tmp_path)
    verified = verify_final_receipt(
        resume_path=resume,
        receipt_path=receipt_path,
        expected_receipt_digest=expected_digest,
        config_path=tmp_path / "config.json",
    )
    assert verified == receipt


def test_forged_nested_draft_digest_is_rejected_even_with_new_receipt_digest(tmp_path):
    resume, receipt_path, receipt, _ = _authorized_artifacts(tmp_path)
    receipt["authorization_report"]["votes"][1]["draft_digest"] = "b" * 64
    receipt["authorization_digest"] = canonical_digest(receipt["authorization_report"])
    receipt_path.write_text(_canonical_json(receipt), encoding="utf-8")
    with pytest.raises(FinalReceiptVerificationError) as failure:
        verify_final_receipt(
            resume_path=resume,
            receipt_path=receipt_path,
            expected_receipt_digest=canonical_digest(receipt),
            config_path=tmp_path / "config.json",
        )
    assert failure.value.code == "AUTHORIZATION_VOTES_MISMATCH"


def test_stale_caller_digest_is_rejected(tmp_path):
    resume, receipt_path, _, expected_digest = _authorized_artifacts(tmp_path)
    stale = ("0" if expected_digest[0] != "0" else "1") + expected_digest[1:]
    with pytest.raises(FinalReceiptVerificationError) as failure:
        verify_final_receipt(
            resume_path=resume,
            receipt_path=receipt_path,
            expected_receipt_digest=stale,
            config_path=tmp_path / "config.json",
        )
    assert failure.value.code == "RECEIPT_DIGEST_MISMATCH"


def test_symlink_receipt_is_rejected(tmp_path):
    resume, receipt_path, _, expected_digest = _authorized_artifacts(tmp_path)
    real_receipt = tmp_path / "stored-receipt.json"
    receipt_path.replace(real_receipt)
    receipt_path.symlink_to(real_receipt)
    with pytest.raises(FinalReceiptVerificationError) as failure:
        verify_final_receipt(
            resume_path=resume,
            receipt_path=receipt_path,
            expected_receipt_digest=expected_digest,
            config_path=tmp_path / "config.json",
        )
    assert failure.value.code == "RECEIPT_ARTIFACT"


def test_copied_receipt_beside_different_job_description_is_rejected(tmp_path):
    resume, receipt_path, _, expected_digest = _authorized_artifacts(tmp_path)
    copied = tmp_path / "copied-application"
    copied.mkdir()
    copied_resume = copied / "resume.md"
    copied_receipt = copied / receipt_path.name
    shutil.copy2(resume, copied_resume)
    shutil.copy2(receipt_path, copied_receipt)
    (copied / "job_description.txt").write_text(
        "A different synthetic role.", encoding="utf-8"
    )

    with pytest.raises(FinalReceiptVerificationError) as failure:
        verify_final_receipt(
            resume_path=copied_resume,
            receipt_path=copied_receipt,
            expected_receipt_digest=expected_digest,
            config_path=tmp_path / "config.json",
        )
    assert failure.value.code == "JOB_DESCRIPTION_DIGEST_MISMATCH"


def test_receipt_under_different_configured_master_is_rejected(tmp_path):
    resume, receipt_path, _, expected_digest = _authorized_artifacts(tmp_path)
    other_master = tmp_path / "other-master.md"
    other_master.write_text("A different trusted master.\n", encoding="utf-8")
    other_config = tmp_path / "other-config.json"
    other_config.write_text(
        json.dumps({"master_resume_path": str(other_master)}), encoding="utf-8"
    )

    with pytest.raises(FinalReceiptVerificationError) as failure:
        verify_final_receipt(
            resume_path=resume,
            receipt_path=receipt_path,
            expected_receipt_digest=expected_digest,
            config_path=other_config,
        )
    assert failure.value.code == "SOURCE_DIGEST_MISMATCH"


def test_candidate_fit_report_digest_tamper_is_rejected(tmp_path):
    resume, receipt_path, receipt, _ = _authorized_artifacts(tmp_path)
    receipt["candidate_fit_report_digest"] = "f" * 64
    receipt_path.write_text(_canonical_json(receipt), encoding="utf-8")
    with pytest.raises(FinalReceiptVerificationError) as failure:
        verify_final_receipt(
            resume_path=resume,
            receipt_path=receipt_path,
            expected_receipt_digest=canonical_digest(receipt),
            config_path=tmp_path / "config.json",
        )
    assert failure.value.code == "CANDIDATE_FIT_REPORT_MISMATCH"


def test_redigested_forged_candidate_fit_pass_is_recomputed_and_rejected(tmp_path):
    resume, receipt_path, receipt, _ = _authorized_artifacts(tmp_path)
    master_text = resume.read_text(encoding="utf-8")
    rejecting_report = assess_candidate_fit(
        master_text,
        KNOCKOUT_JOB_DESCRIPTION,
        run_id=receipt["run_id"],
        case_id=receipt["case_id"],
        as_of_date="2026-07-19",
    )
    assert rejecting_report["passed"] is False
    assert rejecting_report["hard_knockouts"]

    forged_report = json.loads(json.dumps(rejecting_report))
    forged_report["score"] = 100.0
    forged_report["component_scores"] = {
        key: 100.0 for key in forged_report["component_scores"]
    }
    forged_report["hard_knockouts"] = []
    forged_report["extraction_trustworthy"] = True
    forged_report["passed"] = True
    forged_report["recommendation"] = "PROCEED"
    forged_report["codes"] = []

    (tmp_path / "job_description.txt").write_text(
        KNOCKOUT_JOB_DESCRIPTION, encoding="utf-8"
    )
    receipt["job_description_digest"] = canonical_digest(
        KNOCKOUT_JOB_DESCRIPTION
    )
    receipt["candidate_fit_report"] = forged_report
    receipt["candidate_fit_report_digest"] = canonical_digest(forged_report)
    receipt_path.write_text(_canonical_json(receipt), encoding="utf-8")

    with pytest.raises(FinalReceiptVerificationError) as failure:
        verify_final_receipt(
            resume_path=resume,
            receipt_path=receipt_path,
            expected_receipt_digest=canonical_digest(receipt),
            config_path=tmp_path / "config.json",
        )
    assert failure.value.code == "CANDIDATE_FIT_REPORT_MISMATCH"


def test_future_candidate_fit_as_of_date_is_rejected_at_final_boundary(tmp_path):
    resume, receipt_path, receipt, _ = _authorized_artifacts(tmp_path)
    receipt["candidate_fit_report"]["as_of_date"] = "2099-07-19"
    receipt["candidate_fit_report_digest"] = canonical_digest(
        receipt["candidate_fit_report"]
    )
    receipt_path.write_text(_canonical_json(receipt), encoding="utf-8")

    with pytest.raises(FinalReceiptVerificationError) as failure:
        verify_final_receipt(
            resume_path=resume,
            receipt_path=receipt_path,
            expected_receipt_digest=canonical_digest(receipt),
            config_path=tmp_path / "config.json",
        )
    assert failure.value.code == "CANDIDATE_FIT_REPORT_MISMATCH"


def test_candidate_fit_cross_run_replay_is_rejected(tmp_path):
    resume, receipt_path, receipt, _ = _authorized_artifacts(tmp_path)
    receipt["candidate_fit_report"]["run_id"] = "different-run"
    receipt["candidate_fit_report_digest"] = canonical_digest(
        receipt["candidate_fit_report"]
    )
    receipt_path.write_text(_canonical_json(receipt), encoding="utf-8")
    with pytest.raises(FinalReceiptVerificationError) as failure:
        verify_final_receipt(
            resume_path=resume,
            receipt_path=receipt_path,
            expected_receipt_digest=canonical_digest(receipt),
            config_path=tmp_path / "config.json",
        )
    assert failure.value.code == "CANDIDATE_FIT_REPORT_MISMATCH"


def test_legacy_v1_receipt_fails_closed(tmp_path):
    resume, receipt_path, receipt, _ = _authorized_artifacts(tmp_path)
    receipt["schema_version"] = "resume-team-final-receipt/v1"
    receipt_path.write_text(_canonical_json(receipt), encoding="utf-8")
    with pytest.raises(FinalReceiptVerificationError) as failure:
        verify_final_receipt(
            resume_path=resume,
            receipt_path=receipt_path,
            expected_receipt_digest=canonical_digest(receipt),
            config_path=tmp_path / "config.json",
        )
    assert failure.value.code == "RECEIPT_SCHEMA"


def test_cross_host_researcher_identity_is_rejected(tmp_path):
    resume, receipt_path, receipt, _ = _authorized_artifacts(tmp_path)
    receipt["researcher_agent_id"] = "claude:researcher-guard"
    receipt_path.write_text(_canonical_json(receipt), encoding="utf-8")
    with pytest.raises(FinalReceiptVerificationError) as failure:
        verify_final_receipt(
            resume_path=resume,
            receipt_path=receipt_path,
            expected_receipt_digest=canonical_digest(receipt),
            config_path=tmp_path / "config.json",
        )
    assert failure.value.code == "RESEARCHER_ATTESTATION_MISMATCH"


def test_api_host_researcher_and_auditor_identity_is_accepted(tmp_path):
    """The hosted product runtime ("api") is a first-class native host: a
    receipt with distinct api: researcher/auditor identities must verify
    cleanly through both the public schema and the manual cross-check."""
    resume, receipt_path, receipt, _ = _authorized_artifacts(tmp_path)
    receipt["researcher_agent_id"] = "api:researcher-guard"
    receipt["auditor_attestation"]["agent_id"] = "api:auditor-guard"
    receipt_path.write_text(_canonical_json(receipt), encoding="utf-8")
    verified = verify_final_receipt(
        resume_path=resume,
        receipt_path=receipt_path,
        expected_receipt_digest=canonical_digest(receipt),
        config_path=tmp_path / "config.json",
    )
    assert verified == receipt


def test_unrecognized_host_researcher_identity_is_rejected_by_schema(tmp_path):
    """The host enum stays closed: an unrecognized "openai:" prefix must
    still fail closed, at the public schema boundary."""
    resume, receipt_path, receipt, _ = _authorized_artifacts(tmp_path)
    receipt["researcher_agent_id"] = "openai:researcher-guard"
    receipt_path.write_text(_canonical_json(receipt), encoding="utf-8")
    with pytest.raises(FinalReceiptVerificationError) as failure:
        verify_final_receipt(
            resume_path=resume,
            receipt_path=receipt_path,
            expected_receipt_digest=canonical_digest(receipt),
            config_path=tmp_path / "config.json",
        )
    assert failure.value.code == "RECEIPT_SCHEMA"


def test_authorized_resume_docx_entrypoint_passes_bound_digest(tmp_path, monkeypatch):
    resume, receipt_path, receipt, expected_digest = _authorized_artifacts(tmp_path)
    observed = {}

    def fake_create(md_path, output_path, **kwargs):
        observed.update(md_path=md_path, output_path=output_path, kwargs=kwargs)
        return output_path

    monkeypatch.setattr(docx_generator, "_create_resume_from_verified_md", fake_create)
    output = tmp_path / "resume.docx"
    result = docx_generator.create_resume_from_md_authorized(
        str(resume),
        str(output),
        receipt_path=str(receipt_path),
        expected_receipt_digest=expected_digest,
        config_path=str(tmp_path / "config.json"),
    )
    assert result == str(output)
    assert observed["kwargs"]["authorized_draft_digest"] == receipt["draft_digest"]

    with pytest.raises(FinalReceiptVerificationError):
        docx_generator.create_resume_from_md_authorized(
            str(resume),
            str(output),
            receipt_path=str(receipt_path),
            expected_receipt_digest="f" * 64,
            config_path=str(tmp_path / "config.json"),
        )


def test_authorized_resume_docx_roundtrip_and_stale_draft_cannot_replace_output(
    tmp_path,
):
    from resume_integrity_audit import audit_resume_files

    canonical_text = _canonical_resume()
    master = tmp_path / "master.md"
    master.write_text(canonical_text, encoding="utf-8")
    resume, receipt_path, _, expected_digest = _authorized_artifacts(
        tmp_path, canonical_text, master_path=master
    )
    output = tmp_path / "resume.docx"

    result = docx_generator.create_resume_from_md_authorized(
        str(resume),
        str(output),
        receipt_path=str(receipt_path),
        expected_receipt_digest=expected_digest,
        master_path=str(master),
        config_path=str(tmp_path / "config.json"),
    )

    assert result == str(output)
    assert output.is_file()
    assert audit_resume_files(tailored_path=output, master_path=master)["passed"]

    original = output.read_bytes()
    resume.write_text(_canonical_resume(title="Changed Keeper"), encoding="utf-8")
    with pytest.raises(FinalReceiptVerificationError) as failure:
        docx_generator.create_resume_from_md_authorized(
            str(resume),
            str(output),
            receipt_path=str(receipt_path),
            expected_receipt_digest=expected_digest,
            master_path=str(master),
            config_path=str(tmp_path / "config.json"),
        )
    assert failure.value.code == "RESUME_DIGEST_MISMATCH"
    assert output.read_bytes() == original
    assert not list(tmp_path.glob("*.tmp.docx"))


def test_authorized_star_bullet_survives_docx_roundtrip(tmp_path):
    canonical_text = _canonical_resume().replace(
        "• Reviewed signal records and wrote plain findings.",
        "* Reviewed signal records and wrote plain findings.",
    )
    master = tmp_path / "master.md"
    master.write_text(canonical_text, encoding="utf-8")
    resume, receipt_path, _, expected_digest = _authorized_artifacts(
        tmp_path, canonical_text, master_path=master
    )
    output = tmp_path / "resume.docx"

    docx_generator.create_resume_from_md_authorized(
        str(resume),
        str(output),
        receipt_path=str(receipt_path),
        expected_receipt_digest=expected_digest,
        master_path=str(master),
        config_path=str(tmp_path / "config.json"),
    )

    rendered = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    assert "Reviewed signal records and wrote plain findings." in rendered


@pytest.mark.parametrize("supplied_digest", [None, "a" * 64])
def test_direct_resume_docx_entrypoint_fails_before_io(tmp_path, supplied_digest):
    output = tmp_path / "resume.docx"

    with pytest.raises(NativeResumeTeamRequiredError) as failure:
        docx_generator.create_resume_from_md(
            str(tmp_path / "missing-input.md"),
            str(output),
            _authorized_draft_digest=supplied_digest,
        )

    assert failure.value.code == NATIVE_RESUME_TEAM_REQUIRED
    assert not output.exists()
    assert not list(tmp_path.iterdir())


def test_docx_generator_script_fails_closed_without_demo_output(tmp_path):
    completed = subprocess.run(
        [sys.executable, str(ROOT / "docx_generator.py")],
        cwd=tmp_path,
        text=True,
        capture_output=True,
        check=False,
    )

    assert completed.returncode != 0
    assert NATIVE_RESUME_TEAM_REQUIRED in completed.stderr
    assert "Traceback" not in completed.stderr
    assert not (tmp_path / "test_resume.docx").exists()
    assert not list(tmp_path.iterdir())


def test_current_configured_markdown_master_cannot_bypass_docx_authorization(tmp_path):
    from resume_integrity_audit import resolve_master_path

    config = ROOT / "config.json"
    if not config.is_file():
        pytest.skip("repository has no configured master")
    master = Path(resolve_master_path(config))
    if not master.is_file() or master.suffix.lower() not in {".md", ".txt"}:
        pytest.skip("configured master is not a Markdown/text fixture")

    output = tmp_path / "configured-master.docx"
    with pytest.raises(NativeResumeTeamRequiredError) as failure:
        docx_generator.create_resume_from_md(
            str(master),
            str(output),
            master_path=str(master),
            config_path=str(config),
        )
    assert failure.value.code == NATIVE_RESUME_TEAM_REQUIRED
    assert not output.exists()


def test_authorized_cover_letter_entrypoint_revalidates_resume_first(
    tmp_path, monkeypatch
):
    resume, receipt_path, _, expected_digest = _authorized_artifacts(tmp_path)
    cover = tmp_path / "cover_letter.md"
    cover.write_text("cover", encoding="utf-8")
    parsed = {
        "name": "Candidate",
        "contact_info": {},
        "date": "",
        "recipient_info": {},
        "job_title": "",
        "paragraphs": ["Plain paragraph."],
        "closing": "Sincerely,",
    }
    calls = []
    monkeypatch.setattr(docx_generator, "parse_cover_letter_markdown", lambda text: parsed)

    def fake_create_cover(**kwargs):
        calls.append(kwargs)
        document = Document()
        document.add_paragraph(kwargs["name"])
        document.add_paragraph(kwargs["job_title"])
        for paragraph in kwargs["paragraphs"]:
            document.add_paragraph(paragraph)
        document.add_paragraph(kwargs["closing"])
        document.save(kwargs["output_path"])
        return kwargs["output_path"]

    monkeypatch.setattr(
        docx_generator, "create_ats_cover_letter", fake_create_cover
    )
    output = tmp_path / "cover.docx"
    result = docx_generator.create_cover_letter_from_md_authorized(
        str(cover),
        str(output),
        authorized_resume_path=str(resume),
        receipt_path=str(receipt_path),
        expected_receipt_digest=expected_digest,
        job_title="Reviewer",
        company="Example",
        config_path=str(tmp_path / "config.json"),
    )
    assert result == str(output)
    assert calls[0]["job_title"] == "Reviewer"
    assert calls[0]["recipient_info"]["company"] == "Example"

    with pytest.raises(FinalReceiptVerificationError):
        docx_generator.create_cover_letter_from_md_authorized(
            str(cover),
            str(output),
            authorized_resume_path=str(resume),
            receipt_path=str(receipt_path),
            expected_receipt_digest="f" * 64,
            config_path=str(tmp_path / "config.json"),
        )
    assert len(calls) == 1


def test_authorized_cover_formatter_failure_preserves_existing_bytes(
    tmp_path, monkeypatch
):
    resume, receipt_path, _, expected_digest = _authorized_artifacts(tmp_path)
    cover = tmp_path / "cover_letter.md"
    cover.write_text("Plain cover text.", encoding="utf-8")
    output = tmp_path / "cover.docx"
    original = b"existing-cover-letter-bytes"
    output.write_bytes(original)
    parsed = {
        "name": "Candidate",
        "contact_info": {},
        "date": "",
        "recipient_info": {"company": "Example"},
        "job_title": "Reviewer",
        "paragraphs": ["Plain paragraph."],
        "closing": "Sincerely,",
    }
    monkeypatch.setattr(
        docx_generator, "parse_cover_letter_markdown", lambda text: parsed
    )

    def fail_after_partial_write(**kwargs):
        Path(kwargs["output_path"]).write_bytes(b"partial")
        raise RuntimeError("synthetic formatter failure")

    monkeypatch.setattr(
        docx_generator, "create_ats_cover_letter", fail_after_partial_write
    )
    with pytest.raises(RuntimeError, match="synthetic formatter failure"):
        docx_generator.create_cover_letter_from_md_authorized(
            str(cover),
            str(output),
            authorized_resume_path=str(resume),
            receipt_path=str(receipt_path),
            expected_receipt_digest=expected_digest,
            config_path=str(tmp_path / "config.json"),
        )

    assert output.read_bytes() == original
    assert not list(tmp_path.glob(".cover.docx.*.tmp.docx"))


def test_authorized_tracker_false_and_exception_propagate(tmp_path, monkeypatch):
    resume, receipt_path, _, expected_digest = _authorized_artifacts(tmp_path)
    monkeypatch.setattr(tracker_utils, "add_application", lambda **kwargs: False)
    with pytest.raises(tracker_utils.TrackerUpdateError) as failure:
        tracker_utils.add_application_authorized(
            "Example",
            "Reviewer",
            authorized_resume_path=str(resume),
            receipt_path=str(receipt_path),
            expected_receipt_digest=expected_digest,
            config_path=str(tmp_path / "config.json"),
        )
    assert str(failure.value) == "TRACKER_UPDATE_FAILED"

    def explode(**kwargs):
        raise OSError("synthetic")

    monkeypatch.setattr(tracker_utils, "add_application", explode)
    with pytest.raises(tracker_utils.TrackerUpdateError) as failure:
        tracker_utils.add_application_authorized(
            "Example",
            "Reviewer",
            authorized_resume_path=str(resume),
            receipt_path=str(receipt_path),
            expected_receipt_digest=expected_digest,
            config_path=str(tmp_path / "config.json"),
        )
    assert str(failure.value) == "TRACKER_UPDATE_EXCEPTION"


def test_tracker_formatter_failure_preserves_existing_workbook_bytes(
    tmp_path, monkeypatch
):
    tracker = tmp_path / "tracker.xlsx"
    existing = tracker_utils.pd.DataFrame(
        [
            {
                "Company": "Existing",
                "Job Title": "Reviewer",
                "Application Date": "2026-01-01",
            }
        ]
    )
    with tracker_utils.pd.ExcelWriter(tracker, engine="openpyxl") as writer:
        existing.to_excel(writer, sheet_name="Applications", index=False)
    original = tracker.read_bytes()
    monkeypatch.setattr(tracker_utils, "TRACKER_PATH", tracker)

    def fail_format(*args, **kwargs):
        raise RuntimeError("synthetic formatter failure")

    monkeypatch.setattr(tracker_utils, "format_excel_worksheet", fail_format)
    assert tracker_utils.add_application("New", "Analyst") is False
    assert tracker.read_bytes() == original
    assert not list(tmp_path.glob(".tracker.xlsx.*.tmp.xlsx"))
