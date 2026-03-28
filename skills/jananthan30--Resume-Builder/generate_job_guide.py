"""Generate Best-Fit Job Positions Guide as DOCX."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('BEST-FIT JOB POSITIONS GUIDE')
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = RGBColor(31, 78, 121)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Generated from your master resume and scoring profile')
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = RGBColor(89, 89, 89)

doc.add_paragraph()

intro = doc.add_paragraph()
run = intro.add_run('Profile Summary: ')
run.font.bold = True
intro.add_run(
    'This guide identifies the best-fit job positions based on your resume profile, '
    'therapeutic area expertise, credentials, and experience level. Each role is scored '
    'for fit and ranked by priority tier.'
)

intro2 = doc.add_paragraph()
run = intro2.add_run('Key Differentiators: ')
run.font.bold = True
intro2.add_run(
    'Customize the role list below based on your own credentials, therapeutic areas, '
    'and career goals. This template covers common pharma/biotech physician roles.'
)

doc.add_paragraph()


def add_role(rank, title, salary, fit_score, fit_color, why_fit, description,
             requirements, where_to_find, companies, career_path):
    h = doc.add_paragraph()
    run = h.add_run(f'{rank}. {title}')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)

    info = doc.add_paragraph()
    run = info.add_run('Salary Range: ')
    run.font.bold = True
    info.add_run(salary)

    fit = doc.add_paragraph()
    run = fit.add_run('Fit Score: ')
    run.font.bold = True
    run2 = fit.add_run(fit_score)
    run2.font.bold = True
    run2.font.color.rgb = fit_color

    wf = doc.add_paragraph()
    run = wf.add_run('Why This Fits Your Profile: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    wf.add_run(why_fit)

    jd = doc.add_paragraph()
    run = jd.add_run('Job Description: ')
    run.font.bold = True
    jd.add_run(description)

    req = doc.add_paragraph()
    run = req.add_run('Typical Requirements:')
    run.font.bold = True
    for r in requirements:
        doc.add_paragraph(r, style='List Bullet')

    wh = doc.add_paragraph()
    run = wh.add_run('Where to Find These Roles: ')
    run.font.bold = True
    wh.add_run(where_to_find)

    sc = doc.add_paragraph()
    run = sc.add_run('Target Companies: ')
    run.font.bold = True
    sc.add_run(companies)

    cp = doc.add_paragraph()
    run = cp.add_run('Career Path (10-15 Year): ')
    run.font.bold = True
    cp.add_run(career_path)

    doc.add_paragraph()


GREEN = RGBColor(0, 128, 0)
AMBER = RGBColor(204, 153, 0)
GRAY = RGBColor(130, 130, 130)

# ── ROLE 1 ──
add_role(
    1, 'MEDICAL MONITOR / MEDICAL MONITOR PHYSICIAN',
    '$150,000 - $250,000',
    'STRONG FIT (85+)', GREEN,
    'MD REQUIRED for this role \u2014 strongest differentiator for physician candidates. '
    'Clinical training combined with Phase III trial experience (AE/SAE documentation, '
    'safety monitoring, PI point-of-contact) maps directly. MPH adds data-analysis '
    'capabilities. Publications demonstrate scientific rigor expected of medical monitors.',
    'Medical Monitors provide medical and scientific oversight for clinical trials. '
    'They review safety data (adverse events, SAEs, SUSARs), make causality '
    'assessments, advise on eligibility criteria, participate in Data Safety '
    'Monitoring Boards (DSMBs), review protocols and Investigator Brochures, and '
    'serve as the physician voice in drug development \u2014 interpreting clinical data '
    'through a medical lens that non-MDs cannot provide.',
    ['MD or DO degree (required)',
     'Clinical practice experience (2+ years preferred)',
     'Knowledge of ICH-GCP, FDA regulations',
     'Experience with safety data review, AE/SAE assessment',
     'Understanding of clinical trial methodology',
     'Board certification preferred but not always required',
     'Therapeutic area expertise is a plus'],
    'LinkedIn ("Medical Monitor" OR "Medical Monitor Physician"), ClinicalCrossing, '
    'Indeed, Glassdoor, company career pages',
    'Medpace, ICON, Syneos Health, Parexel, PPD/Thermo Fisher, IQVIA, Covance/Labcorp, '
    'Fortrea, Worldwide Clinical Trials, Novartis, Pfizer, Merck, J&J, AbbVie, Regeneron, '
    'Genmab, Anavex, smaller biotechs',
    'Medical Monitor \u2192 Senior Medical Monitor \u2192 Medical Director \u2192 '
    'VP Clinical Development \u2192 CMO'
)

# ── ROLE 2 ──
add_role(
    2, 'DRUG SAFETY PHYSICIAN / PHARMACOVIGILANCE PHYSICIAN',
    '$150,000 - $220,000',
    'STRONG FIT (80+)', GREEN,
    'MD REQUIRED. Clinical training provides medical judgment for causality '
    'assessments and benefit-risk evaluations. Phase III trial AE/SAE review '
    'experience maps directly. Publications on drug safety topics demonstrate '
    'pharmacovigilance thinking. Critical care experience translates to safety '
    'signal detection.',
    'Drug Safety Physicians evaluate Individual Case Safety Reports (ICSRs), perform '
    'causality assessments, contribute to aggregate reports (PSURs/PBRERs, DSURs), '
    'participate in signal detection and risk management planning, review REMS, and '
    'provide medical input on safety-related regulatory submissions.',
    ['MD or DO degree (required)',
     '2+ years clinical experience',
     'Knowledge of PV regulations (ICH E2A-E2F)',
     'Experience with MedDRA coding, ICSRs',
     'Understanding of aggregate safety reporting',
     'Familiarity with safety databases (Argus, AERS)',
     'Strong medical writing skills'],
    'LinkedIn ("Drug Safety Physician" OR "PV Physician" OR "Safety Physician"), '
    'PharmaCrossing, company career pages',
    'Pfizer, Merck, J&J, AbbVie, Novartis, Roche, AstraZeneca, GSK, Sanofi, BMS, '
    'IQVIA PV division, ProPharma, Celerion, Worldwide Clinical Trials',
    'Drug Safety Physician \u2192 Senior DSP \u2192 Head of Drug Safety \u2192 '
    'VP Pharmacovigilance \u2192 Chief Safety Officer'
)

# ── ROLE 3 ──
add_role(
    3, 'ASSOCIATE MEDICAL DIRECTOR / MEDICAL DIRECTOR (CLINICAL DEVELOPMENT)',
    '$180,000 - $300,000',
    'STRONG FIT (75+)', GREEN,
    'MD REQUIRED. This is the top-tier pharma role for physicians. Clinical '
    'trial experience, peer-reviewed publications, and an MPH create a compelling '
    'profile. Physicians provide the clinical-scientific leadership that non-MD '
    'candidates cannot. Gap: most Medical Directors have 3-5+ years in pharma or '
    'completed a residency/fellowship.',
    'Medical Directors lead clinical development programs from a medical-scientific '
    'perspective. They design protocols, interpret clinical data, chair safety '
    'committees, engage with KOLs and regulatory agencies, author clinical sections '
    'of regulatory submissions, and serve as the medical expert for all trial-related '
    'decisions. They are the physician-leader of the clinical program.',
    ['MD or DO degree (required)',
     'Board certification preferred',
     '3-5+ years pharma/biotech or clinical research experience',
     'Therapeutic area expertise',
     'Protocol design and clinical document authoring',
     'Regulatory interaction experience (FDA, EMA)',
     'KOL engagement and advisory board management'],
    'LinkedIn, BioSpace, company career pages. Search "Medical Director" OR '
    '"Associate Medical Director" AND "Clinical Development"',
    'All major pharma (Pfizer, Merck, J&J, Novartis, Roche, AstraZeneca, Sanofi, '
    'BMS, Gilead, Amgen, Regeneron), mid-size biotech, and CROs with medical '
    'monitoring divisions',
    'Associate Medical Director \u2192 Medical Director \u2192 Senior Medical Director '
    '\u2192 VP Clinical Development \u2192 SVP \u2192 CMO'
)

# ── ROLE 4 ──
add_role(
    4, 'CLINICAL SCIENTIST / ASSOCIATE DIRECTOR CLINICAL SCIENCE',
    '$130,000 - $200,000',
    'STRONG FIT (75+)', GREEN,
    'MD provides clinical judgment for protocol design and data interpretation. '
    'Publications demonstrate scientific rigor and writing capability. Phase III '
    'trial experience (medical data review, safety trend analysis) maps directly. '
    'MPH with data/informatics focus adds depth. MD + publications + trial '
    'operations is a rare and valuable combination.',
    'Clinical Scientists contribute to clinical development strategy, protocol '
    'design, and medical data review. They author/review clinical documents '
    '(protocols, IBs, CSRs, regulatory submissions), perform ongoing medical data '
    'review and safety data interpretation, support site selection, prepare '
    'presentations for Investigator Meetings, and contribute to publications strategy.',
    ['PhD, PharmD, MD, or MS in life sciences',
     '5+ years in clinical research or drug development',
     'Protocol development and clinical document authoring',
     'Medical data review and interpretation',
     'Therapeutic area expertise preferred',
     'Strong scientific writing',
     'Phase I-III trial experience'],
    'LinkedIn, Glassdoor, BioSpace. Search "Clinical Scientist" OR '
    '"Clinical Research Scientist" OR "Associate Director Clinical Science"',
    'Genmab, Regeneron, Celcuity, Recursion, BMS, Gilead, Amgen, Sanofi, Takeda, '
    'Astellas, Jazz Pharmaceuticals, Blueprint Medicines, smaller biotechs',
    'Clinical Scientist \u2192 Senior CS \u2192 Associate Director \u2192 '
    'Director Clinical Science \u2192 VP Clinical Development'
)

# ── ROLE 5 ──
add_role(
    5, 'CLINICAL OPERATIONS MANAGER / ASSOCIATE DIRECTOR',
    '$130,000 - $180,000',
    'MODERATE-STRONG FIT (70+)', AMBER,
    'Clinical site experience is the proof point: concurrent Phase III trials, '
    'CRO oversight, vendor management, TMF completeness, study start-up, '
    'cross-functional coordination. CRO-side experience adds perspective. '
    'MD provides clinical credibility most Ops professionals lack. Gap: most '
    'Directors have 8-10+ years specifically in operations.',
    'Clinical Operations leaders oversee planning, execution, and delivery of '
    'clinical trials. They manage CRO relationships, drive study start-up, oversee '
    'site management, ensure TMF completeness and inspection readiness, manage '
    'budgets and timelines, coordinate cross-functional teams, and develop SOPs.',
    ['Bachelor minimum (advanced degree preferred)',
     '8+ years in clinical operations or research',
     'CRO management and vendor oversight',
     'Study start-up, site management, TMF',
     'ICH-GCP, FDA/EMA regulations',
     'Budget management and contract negotiation',
     'Matrixed environment experience'],
    'LinkedIn, BioSpace, Indeed. Search "Clinical Operations" AND '
    '("Associate Director" OR "Director" OR "Manager")',
    'Anavex, Genmab, Regeneron, Sanofi, Pfizer, Merck, J&J, AbbVie, Medpace, '
    'ICON, Syneos, Parexel, PPD/Thermo Fisher, IQVIA, Fortrea, mid-size biotechs',
    'Manager \u2192 Associate Director \u2192 Director \u2192 Senior Director '
    '\u2192 VP Clinical Operations \u2192 COO'
)

# ── ROLE 6 ──
add_role(
    6, 'MEDICAL SCIENCE LIAISON (MSL)',
    '$140,000 - $200,000 + bonus',
    'MODERATE-STRONG FIT (70+)', AMBER,
    'MD is the gold standard for MSL roles (most MSLs have PharmD/PhD \u2014 an MD '
    'stands out). Publications demonstrate scientific credibility for KOL '
    'engagement. Clinical experience provides real-world perspective KOLs respect. '
    'Therapeutic area expertise opens disease-specific MSL roles. '
    'Gap: no prior MSL/medical affairs experience, and travel is 60-80%.',
    'MSLs are field-based scientific experts who build KOL relationships, present '
    'clinical data at conferences, respond to unsolicited medical inquiries, train '
    'internal teams, gather field medical insights, and support Investigator-Initiated '
    'Studies (IIS). MSLs do NOT sell \u2014 they engage in peer-to-peer scientific exchange.',
    ['Advanced degree: MD, PharmD, PhD, or DO',
     'Therapeutic area expertise matching pipeline',
     '2+ years clinical or research experience',
     'Strong presentation and communication',
     'Publication record preferred',
     'Willingness to travel 60-80%',
     'Prior MSL experience preferred (not required for MDs)'],
    'LinkedIn ("Medical Science Liaison" + therapeutic area), MSL Society job board, '
    'MedReps, BioSpace',
    'Pfizer, Merck, J&J, AbbVie, Gilead (HIV MSL), BMS, Novartis, Sanofi, Regeneron, '
    'Amgen, AstraZeneca, GSK, Takeda, Jazz Pharmaceuticals',
    'MSL \u2192 Senior MSL \u2192 MSL Director \u2192 National MSL Director '
    '\u2192 VP Medical Affairs \u2192 CMO'
)

# ── ROLE 7 ──
add_role(
    7, 'SENIOR MEDICAL WRITER / ASSOCIATE DIRECTOR MEDICAL WRITING',
    '$110,000 - $170,000',
    'MODERATE FIT (65+)', AMBER,
    'Peer-reviewed publications and book chapters prove scientific writing '
    'ability. MD provides clinical interpretation skills. Phase III experience '
    'gives understanding of CSRs, protocols, and regulatory documents. '
    'Gap: no formal medical writing in pharma/CRO. Publications are your portfolio.',
    'Medical Writers create regulatory and scientific documents: Clinical Study '
    'Reports (CSRs), protocols, Investigator Brochures (IBs), safety narratives, '
    'regulatory submissions (IND, NDA, BLA), briefing documents, and publications. '
    'They translate complex clinical data into clear, compliant documents.',
    ['Advanced degree (PhD, PharmD, MD, MS)',
     '3-5+ years medical writing or scientific writing',
     'ICH guidelines and regulatory document knowledge',
     'Document management systems (Veeva Vault)',
     'Clinical trial methodology understanding',
     'AMWA/EMWA certification is a plus'],
    'LinkedIn, AMWA job board, MedComms Networking, BioSpace. Search "Medical Writer" '
    'OR "Regulatory Writer"',
    'Syneos, ICON, Parexel, IQVIA, PPD/Thermo Fisher, Certara, ProPharma, '
    'Cactus Communications, pharma in-house writing departments',
    'Medical Writer \u2192 Senior \u2192 Principal \u2192 Associate Director '
    '\u2192 Director Medical Writing \u2192 VP Regulatory Affairs'
)

# ── ROLE 8 ──
add_role(
    8, 'REAL WORLD EVIDENCE (RWE) SCIENTIST / OUTCOMES RESEARCHER',
    '$120,000 - $170,000',
    'MODERATE FIT (65+)', AMBER,
    'MPH (Health Informatics) + Python/SQL/BigQuery + data science projects '
    '= strong analytical foundation. Clinical experience provides real-world '
    'patient context. Predictive modeling with real-world datasets is directly '
    'RWE-relevant. Gap: no formal HEOR/outcomes research in pharma.',
    'RWE Scientists design and execute real-world evidence studies using EHRs, '
    'claims databases, and registries. They support market access with evidence of '
    'real-world effectiveness, contribute to FDA RWE submissions, perform '
    'comparative effectiveness research, and support product lifecycle management.',
    ['Masters or PhD in epi, biostatistics, health economics, or public health',
     'Experience with real-world data (claims, EHR, registries)',
     'Programming (Python, R, SAS, SQL)',
     'Study design methodology',
     'FDA RWE guidance knowledge preferred',
     'Strong scientific writing'],
    'LinkedIn, ISPOR job board, BioSpace. Search "Real World Evidence" OR "RWE" '
    'OR "HEOR" OR "Outcomes Research"',
    'IQVIA RWE, Evidera, Aetion, Flatiron Health, Optum, CVS Health, '
    'Pfizer, Merck, J&J, Regeneron, Amgen',
    'RWE Scientist \u2192 Senior \u2192 Associate Director HEOR \u2192 '
    'Director RWE \u2192 VP Evidence Generation'
)

# ── ROLE 9 ──
add_role(
    9, 'MEDICAL INFORMATION SPECIALIST / ASSOCIATE DIRECTOR',
    '$100,000 - $160,000',
    'MODERATE FIT (60+)', GRAY,
    'MD provides clinical knowledge for complex medical inquiries. Publications '
    'demonstrate literature synthesis skills. Clinical breadth across specialties '
    'gives wide medical knowledge base. Gap: no formal medical information or '
    'medical affairs experience.',
    'Medical Information professionals respond to medical inquiries from HCPs, '
    'patients, and internal teams. They create standard response documents, perform '
    'medical literature searches, support medical content review, manage AE intake '
    'and escalation, and contribute to scientific communications strategy.',
    ['Advanced degree (MD, PharmD, PhD preferred)',
     'FDA regulations for medical information',
     'Medical literature search and synthesis',
     'Medical information databases',
     'Strong written and verbal communication',
     'Therapeutic area knowledge preferred'],
    'LinkedIn, Indeed. Search "Medical Information" AND ("Specialist" OR "Manager" '
    'OR "Director")',
    'Pfizer, Merck, J&J, AbbVie, Novartis, Roche, BMS, Sanofi, ProPharma, '
    'Inizio Medical, Ashfield/Inizio Engage, ICON',
    'Specialist \u2192 Senior \u2192 Manager \u2192 Associate Director '
    '\u2192 Director Medical Information \u2192 VP Medical Affairs'
)

# ── SUMMARY PAGE ──
doc.add_page_break()
h = doc.add_paragraph()
run = h.add_run('SUMMARY: PRIORITY APPLICATION STRATEGY')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(31, 78, 121)

doc.add_paragraph()

t1 = doc.add_paragraph()
run = t1.add_run('TIER 1 \u2014 APPLY IMMEDIATELY (MD Required = Your Biggest Edge)')
run.font.bold = True
run.font.color.rgb = GREEN
for item in [
    'Medical Monitor / Medical Monitor Physician \u2014 $150-250K \u2014 MD required, highest fit',
    'Drug Safety Physician / PV Physician \u2014 $150-220K \u2014 MD required, safety expertise',
    'Associate Medical Director (Clinical Development) \u2014 $180-300K \u2014 MD required, highest ceiling',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

t2 = doc.add_paragraph()
run = t2.add_run('TIER 2 \u2014 STRONG APPLICATIONS (MD is Differentiator)')
run.font.bold = True
run.font.color.rgb = AMBER
for item in [
    'Clinical Scientist / Assoc. Director Clinical Science \u2014 $130-200K \u2014 publications + trial experience',
    'Clinical Operations Manager / Assoc. Director \u2014 $130-180K \u2014 direct operations experience',
    'Medical Science Liaison \u2014 $140-200K \u2014 MD + publications for KOL engagement',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

t3 = doc.add_paragraph()
run = t3.add_run('TIER 3 \u2014 VIABLE APPLICATIONS (Transferable Skills)')
run.font.bold = True
run.font.color.rgb = GRAY
for item in [
    'Senior Medical Writer \u2014 $110-170K \u2014 publications as portfolio',
    'RWE Scientist / Outcomes Researcher \u2014 $120-170K \u2014 MPH + data skills',
    'Medical Information \u2014 $100-160K \u2014 MD + clinical breadth',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# DO NOT APPLY
dn = doc.add_paragraph()
run = dn.add_run('DO NOT APPLY TO THESE ROLES:')
run.font.bold = True
run.font.color.rgb = RGBColor(200, 0, 0)
for item in [
    'CRA / Senior CRA \u2014 Requires 5+ years of specific site monitoring visit experience you lack',
    'Clinical Research Coordinator \u2014 Below your qualification level; you will leave quickly',
    'Medical Assistant / Clinical Assistant \u2014 Far below MD + MPH level',
    'Entry-level anything \u2014 Wastes your credentials and worsens job-hopping pattern',
    '5+ roles at a single company \u2014 Signals desperation, not targeted interest',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# Key advice
adv = doc.add_paragraph()
run = adv.add_run('KEY JOB SEARCH ADVICE')
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = RGBColor(31, 78, 121)

advice_items = [
    'Focus 70% of applications on Tier 1 roles where MD is REQUIRED \u2014 '
    'least competition from non-physicians, highest salary ceilings.',
    'Apply to 1-2 roles max per company. Applying to 5+ signals desperation.',
    'Your cover letter should lead with your MD and clinical trial experience '
    '\u2014 the physician credential is the headline.',
    'Prestige academic credentials (top MPH, strong GPA) should appear in paragraph 1.',
    'Publications aligned to therapeutic areas are assets for disease-specific roles. '
    'Trial experience in specific TAs opens doors at companies in those spaces.',
    'Target companies with active Phase II-III programs in your TAs: '
    'HIV (Gilead, ViiV/GSK, Merck), CV (Novartis, Bayer, BMS), CNS (Biogen, Sage, Anavex).',
    'Contract/consulting roles (Medical Monitor contractor at $100-150/hr) can bridge '
    'income while searching for permanent positions. Check: Medpace, ProPharma, Premier Research.',
    'AI training reviewer roles ($50-150/hr, remote) are excellent side income \u2014 '
    'leverage your MD + AI/ML experience for platforms like Scale AI, Invisible Technologies.',
]
for a in advice_items:
    doc.add_paragraph(a, style='List Bullet')

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Generated March 12, 2026 | Resume Builder v5.1')
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(150, 150, 150)

doc.save('Best_Fit_Job_Positions_Guide.docx')
print('DOCX created successfully')
