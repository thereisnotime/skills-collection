"""
DOCX Generator for Resume Builder

Supports multiple resume formats:
1. ATS-Optimized (Workday/ATS Compliant)
2. Harvard Style (Harvard OCS Format)

ATS Design Rules (2026 Best Practices):
- NO columns, tables, text boxes, graphics, icons, or headers/footers
- YES to bold text, all-caps headers, bullet points, horizontal lines
- Font: Calibri (ATS-friendly sans-serif, corporate standard)
- Size hierarchy: Name 20pt, Section Headers 11pt bold, Body 10.5pt, Contact 10pt
- Color accent: Dark navy (#1F3864) for name and section headers only
- Company names: Title Case (not ALL CAPS) — only job titles are uppercase
- Line spacing: 1.15 for readability
- Contact info in main body, not header
- Linear single-column layout for easy parsing

Harvard Style Design Rules (Harvard Office of Career Services):
- Font: Times New Roman (traditional academic)
- Name: 14pt bold, centered
- Section Headers: Bold with underline, not all-caps
- Education section comes FIRST (Harvard prioritizes education)
- Dates are RIGHT-ALIGNED using tab stops
- Organization/company names are BOLD
- Job titles in regular weight
- Clean look without horizontal lines
- Margins: 0.5" to 1"
- No graphics or fancy formatting
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# -- Default font and color constants --
FONT_BODY = 'Calibri'
FONT_SIZE_NAME = 20
FONT_SIZE_HEADER = 11
FONT_SIZE_BODY = 10.5
FONT_SIZE_CONTACT = 10
FONT_SIZE_BULLET = 10.5
FONT_SIZE_PUB = 9.5
COLOR_NAVY = RGBColor(0x1F, 0x38, 0x64)   # Dark navy accent (#1F3864)
COLOR_BLACK = RGBColor(0, 0, 0)
COLOR_DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
LINE_SPACING = 1.15

# Acronyms / abbreviations that should stay ALL CAPS in company names
_KEEP_UPPER = {
    'LLC', 'INC', 'LLP', 'LP', 'PLC', 'SA', 'AG', 'CO', 'LTD',
    'USA', 'US', 'UK', 'NY', 'NJ', 'CT', 'CA', 'WI', 'MD', 'DC',
    'CRO', 'EDC', 'ICU', 'IRB', 'FDA', 'EMA', 'GCP', 'ICH',
    'IQVIA', 'ICON', 'IBM', 'NIH', 'CDC', 'WHO', 'HCG', 'MPH',
    'EMR', 'EHR', 'IT', 'AI', 'ML', 'HR', 'QA', 'R&D',
    'DR', 'DR.', 'K.', 'ESC', 'SME', 'CNS', 'EPM',
    'HIV', 'OB/GYN', 'OBGYN', 'NYU', 'MIT', 'UCLA',
}


def smart_title_case(text: str) -> str:
    """
    Convert ALL-CAPS company names to Title Case, preserving known acronyms.

    Examples:
        'YALE SCHOOL OF PUBLIC HEALTH' -> 'Yale School of Public Health'
        'COVANCE/LABCORP DRUG DEVELOPMENT' -> 'Covance/LabCorp Drug Development'
        'SAINT MICHAEL MEDICAL CENTER' -> 'Saint Michael Medical Center'
        'DR. K. ISWARA GASTROENTEROLOGY OFFICE' -> 'Dr. K. Iswara Gastroenterology Office'
    """
    if not text or not text.isupper():
        return text  # only transform ALL-CAPS strings

    # Small words that stay lowercase in title case (unless first word)
    small_words = {'of', 'the', 'and', 'in', 'for', 'at', 'by', 'to', 'a', 'an', 'on'}

    # Handle slash-separated parts (e.g., COVANCE/LABCORP)
    if '/' in text:
        return '/'.join(smart_title_case(part) for part in text.split('/'))

    # Handle hyphenated parts
    if ' - ' in text:
        return ' - '.join(smart_title_case(part) for part in text.split(' - '))

    words = text.split()
    result = []
    for i, word in enumerate(words):
        # Strip trailing punctuation for lookup
        core = word.rstrip('.,;:')
        if core.upper() in _KEEP_UPPER or word.upper() in _KEEP_UPPER:
            result.append(word.upper() if word.isupper() else word)
        elif i > 0 and core.lower() in small_words:
            result.append(word.lower())
        else:
            result.append(word.capitalize())

    return ' '.join(result)


def set_font(run, font_name=None, size=None, bold=False):
    """Set font properties for a run. Defaults to Calibri 10.5pt."""
    font_name = font_name or FONT_BODY
    size = size or FONT_SIZE_BODY
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    # Set font for East Asian text as well
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def _set_line_spacing(p, multiple=None):
    """Set paragraph line spacing as a multiple of single (240-unit) spacing."""
    if multiple is None:
        multiple = LINE_SPACING
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(int(240 * multiple)))
    spacing.set(qn('w:lineRule'), 'auto')


def add_horizontal_line(doc, thick=False):
    """Add a proper hairline paragraph-border rule (replaces underscore text)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8' if thick else '6')   # 1pt thick or 0.75pt hairline
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _setup_heading_style(doc):
    """
    Create a custom 'SectionHeader' style based on Normal (not Heading 1).

    Using a Normal-based style avoids Word's collapse/expand toggle triangle
    that appears on all built-in Heading styles. ATS parsers like Workday
    identify sections via ALL-CAPS text heuristics and bold formatting,
    which this custom style preserves.
    """
    styles = doc.styles
    try:
        # Create custom style based on Normal to avoid toggle triangles
        if 'SectionHeader' not in [s.name for s in styles]:
            from docx.enum.style import WD_STYLE_TYPE
            sh = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
            sh.base_style = styles['Normal']
            sh.font.name = FONT_BODY
            sh.font.size = Pt(FONT_SIZE_HEADER)
            sh.font.bold = True
            sh.font.color.rgb = COLOR_NAVY
            sh.font.italic = False
            sh.paragraph_format.space_before = Pt(14)
            sh.paragraph_format.space_after = Pt(6)
            sh.paragraph_format.keep_with_next = True
            sh.paragraph_format.left_indent = Pt(0)
    except Exception:
        pass  # If style creation fails, add_section_header falls back to Normal


def add_section_header(doc, text):
    """
    Add an all-caps section header using a custom Normal-based style.

    Uses 'SectionHeader' style (based on Normal, not Heading 1) to avoid
    Word's collapse/expand toggle triangle. ATS parsers identify sections
    via ALL-CAPS bold text, which this approach preserves.
    """
    try:
        p = doc.add_paragraph(style='SectionHeader')
    except KeyError:
        p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.left_indent = Pt(0)

    run = p.add_run(text.upper())
    set_font(run, font_name=FONT_BODY, size=FONT_SIZE_HEADER, bold=True)
    run.font.color.rgb = COLOR_NAVY
    run.font.italic = False
    # Add 1pt letter-spacing for all-caps headers
    rPr = run._element.get_or_add_rPr()
    spacing_el = OxmlElement('w:spacing')
    spacing_el.set(qn('w:val'), '20')   # 20 twentieths-of-pt = 1pt
    rPr.append(spacing_el)


def add_bullet_point(doc, text, bold_parts=None):
    """Add a bullet point with optional bold parts."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)
    _set_line_spacing(p)

    if bold_parts:
        # Split text and bold specific parts
        remaining = text
        for bold_text in bold_parts:
            if bold_text in remaining:
                parts = remaining.split(bold_text, 1)
                if parts[0]:
                    run = p.add_run(parts[0])
                    set_font(run, size=FONT_SIZE_BULLET)
                run = p.add_run(bold_text)
                set_font(run, size=FONT_SIZE_BULLET, bold=True)
                remaining = parts[1] if len(parts) > 1 else ''
        if remaining:
            run = p.add_run(remaining)
            set_font(run, size=FONT_SIZE_BULLET)
    else:
        run = p.add_run(text)
        set_font(run, size=FONT_SIZE_BULLET)

    return p


def _render_citation_runs(p, text, prefix=""):
    """
    Render a citation string into runs on paragraph p with:
    - Bold year
    - Italic journal / conference name (the segment between the last '. ' and the year,
      for journal-format citations where the year is followed by ';' or digits)
    Falls back to plain+bold-year for non-standard formats.
    """
    full = prefix + text
    FN = FONT_BODY
    SZ = FONT_SIZE_PUB

    year_m = re.search(r'\b((?:19|20)\d{2})\b', full)
    if not year_m:
        run = p.add_run(full)
        set_font(run, FN, SZ)
        return

    year_str = year_m.group(1)
    y_start = year_m.start()
    before_year = full[:y_start]
    after_year = full[y_start + len(year_str):]

    # Detect journal-article format: year immediately followed by ";" or digit (vol/issue)
    after_stripped = after_year.lstrip()
    is_journal = after_stripped.startswith(';') or (after_stripped and after_stripped[0].isdigit())

    if is_journal:
        # Find journal name: segment from last ". " before year to year (stripped of trailing .,; )
        last_sep = max(before_year.rfind('. '), before_year.rfind('? '))
        if last_sep != -1:
            pre_journal = full[:last_sep + 2]       # includes the ". "
            journal_raw = before_year[last_sep + 2:].rstrip('. ;,')
            post_journal = before_year[last_sep + 2 + len(journal_raw):]  # any ". " after journal
            # Only italicise if it looks like a real journal name (3–80 chars, not a year)
            if 3 <= len(journal_raw) <= 80 and not re.match(r'^\d', journal_raw):
                if pre_journal:
                    r = p.add_run(pre_journal)
                    set_font(r, FN, SZ)
                r = p.add_run(journal_raw)
                set_font(r, FN, SZ)
                r.font.italic = True
                if post_journal:
                    r = p.add_run(post_journal)
                    set_font(r, FN, SZ)
                yr = p.add_run(year_str)
                set_font(yr, FN, SZ, bold=True)
                if after_year:
                    r = p.add_run(after_year)
                    set_font(r, FN, SZ)
                return  # success path

    # Fallback: just bold the year, no journal italic
    if before_year:
        r = p.add_run(before_year)
        set_font(r, FN, SZ)
    yr = p.add_run(year_str)
    set_font(yr, FN, SZ, bold=True)
    if after_year:
        r = p.add_run(after_year)
        set_font(r, FN, SZ)


def add_publication_entry(doc, text, number=None):
    """
    Add a publication citation with hanging indent (academic reference style).
    Long citations wrap cleanly: first line flush-left, continuation indented 0.35".
    Year is bold; journal/conference name is italic for journal-format citations.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.35)  # Hanging indent

    prefix = f"{number}.\u2002" if number else ""  # en-space after period
    _render_citation_runs(p, text, prefix=prefix)
    return p


def parse_markdown_bold(text):
    """
    Parse markdown **bold** syntax and return clean text with bold parts.

    Returns:
        tuple: (clean_text, list_of_bold_parts)
    """
    # Find all **text** patterns
    bold_pattern = r'\*\*([^*]+)\*\*'
    bold_parts = re.findall(bold_pattern, text)

    # Remove the ** markers from text
    clean_text = re.sub(bold_pattern, r'\1', text)

    return clean_text, bold_parts


def extract_metrics(text):
    """Extract metrics and numbers to bold them."""
    # First check for markdown bold syntax
    clean_text, markdown_bold = parse_markdown_bold(text)

    if markdown_bold:
        # If markdown bold was used, return those parts
        return clean_text, markdown_bold

    # Otherwise, auto-detect metrics
    # Find patterns like: 95%, $2.5M, 8 centers, 11,300+, 4 hours, 91% sensitivity
    patterns = [
        r'\d+%',  # Percentages
        r'\$[\d.,]+[MKB]?',  # Money
        r'\d+[\d,]*\+?\s*(?:years?|months?|centers?|sites?|projects?|patients?|teams?|members?)',  # Counts
        r'\d+[\d,]*\+',  # Numbers with plus
    ]

    bold_parts = []
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        bold_parts.extend(matches)

    return text, bold_parts


def create_ats_resume(
    output_path,
    name,
    contact_info,  # dict with city, state, phone, email, linkedin
    summary,
    core_competencies,  # list of keywords
    experience,  # list of dicts with title, company, location, dates, bullets
    education,  # list of dicts with degree, school, location
    certifications,  # list of strings
    professional_memberships=None,  # optional list of strings
    publications=None,  # optional list of strings or dict with categories
    publications_footer=None,  # optional footer line (e.g., Google Scholar link)
    projects=None,  # optional list of dicts with title, dates, bullets
):
    """
    Create an ATS-optimized resume DOCX.

    Args:
        output_path: Path to save the DOCX file
        name: Full name with credentials (e.g., "JANE DOE, M.D.")
        contact_info: dict with keys: city, state, zip, phone, email, linkedin
        summary: Professional summary text (3-4 lines)
        core_competencies: List of keyword strings
        experience: List of job dicts with: title, company, location, dates, bullets
        education: List of education dicts with: degree, school, location, dates
        certifications: List of certification strings
        professional_memberships: Optional list of membership strings
        publications: Optional list of publication strings, or dict with category keys
                      (e.g., {'Peer-Reviewed Journal Articles': [...], 'Book Chapters': [...]})
    """
    doc = Document()

    # Configure Heading 1 style for Workday-compatible section header parsing
    _setup_heading_style(doc)

    # Set margins (0.7" for clean readability)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # ===== NAME (Bold, Centered, 20pt, Navy) =====
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(3)
    name_run = name_para.add_run(name.upper())
    set_font(name_run, size=FONT_SIZE_NAME, bold=True)
    name_run.font.color.rgb = COLOR_NAVY

    # ===== CONTACT INFO (In main body, not header) =====
    # Use pipe separator and filter out empty fields
    SEP = '  |  '
    contact_parts = []
    city_state = f"{contact_info.get('city', '')}, {contact_info.get('state', '')} {contact_info.get('zip', '')}".strip(' ,')
    if city_state.strip():
        contact_parts.append(city_state.strip())
    if contact_info.get('phone', '').strip():
        contact_parts.append(contact_info['phone'].strip())
    if contact_info.get('email', '').strip():
        contact_parts.append(contact_info['email'].strip())

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_before = Pt(0)
    contact_para.paragraph_format.space_after = Pt(0)
    contact_run = contact_para.add_run(SEP.join(contact_parts))
    set_font(contact_run, size=FONT_SIZE_CONTACT)
    contact_run.font.color.rgb = COLOR_DARK_GRAY

    # LinkedIn/Portfolio on second line
    if contact_info.get('linkedin'):
        link_para = doc.add_paragraph()
        link_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        link_para.paragraph_format.space_before = Pt(1)
        link_para.paragraph_format.space_after = Pt(0)
        link_run = link_para.add_run(contact_info['linkedin'])
        set_font(link_run, size=FONT_SIZE_CONTACT)
        link_run.font.color.rgb = COLOR_DARK_GRAY

    # Thicker rule closes the header block
    add_horizontal_line(doc, thick=True)

    # ===== PROFESSIONAL SUMMARY =====
    add_section_header(doc, 'PROFESSIONAL SUMMARY')

    summary_para = doc.add_paragraph()
    summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    summary_para.paragraph_format.space_before = Pt(2)
    summary_para.paragraph_format.space_after = Pt(6)
    _set_line_spacing(summary_para)
    summary_run = summary_para.add_run(summary)
    set_font(summary_run, size=FONT_SIZE_BODY)

    # ===== CORE COMPETENCIES =====
    add_horizontal_line(doc)
    add_section_header(doc, 'CORE COMPETENCIES')

    # Display 3 competencies per row separated by pipe — compact and ATS-safe
    ROW = 3
    for i in range(0, len(core_competencies), ROW):
        row_items = core_competencies[i:i + ROW]
        comp_para = doc.add_paragraph(style='List Bullet')
        comp_para.paragraph_format.space_after = Pt(2)
        comp_para.paragraph_format.left_indent = Inches(0.25)
        for j, item in enumerate(row_items):
            if j > 0:
                sep_run = comp_para.add_run('  |  ')
                set_font(sep_run, size=FONT_SIZE_BODY)
                sep_run.font.color.rgb = COLOR_DARK_GRAY
            comp_run = comp_para.add_run(item)
            set_font(comp_run, size=FONT_SIZE_BODY)

    # ===== PROFESSIONAL EXPERIENCE =====
    add_horizontal_line(doc)
    add_section_header(doc, 'PROFESSIONAL EXPERIENCE')

    for job in experience:
        # Line 1: JOB TITLE (bold, ALL-CAPS, left) .......... Dates (right-aligned)
        title_para = doc.add_paragraph()
        title_para.paragraph_format.space_before = Pt(10)
        title_para.paragraph_format.space_after = Pt(1)
        if job.get('dates'):
            title_para.paragraph_format.tab_stops.add_tab_stop(
                Inches(6.3), WD_TAB_ALIGNMENT.RIGHT
            )
        title_run = title_para.add_run(job['title'].upper())
        set_font(title_run, size=FONT_SIZE_BODY, bold=True)
        if job.get('dates'):
            title_para.add_run('\t')
            dates_run = title_para.add_run(job['dates'])
            set_font(dates_run, size=FONT_SIZE_BODY)
            dates_run.font.color.rgb = COLOR_DARK_GRAY

        # Line 2: Company Name (Title Case, bold), Location (regular)
        company_para = doc.add_paragraph()
        company_para.paragraph_format.space_before = Pt(0)
        company_para.paragraph_format.space_after = Pt(3)
        company_name = smart_title_case(job['company'])
        company_run = company_para.add_run(company_name)
        set_font(company_run, size=FONT_SIZE_BODY, bold=True)
        company_run.font.color.rgb = COLOR_DARK_GRAY
        if job.get('location'):
            loc_text = smart_title_case(job['location']) if job['location'].isupper() else job['location']
            loc_run = company_para.add_run(f", {loc_text}")
            set_font(loc_run, size=FONT_SIZE_BODY)
            loc_run.font.color.rgb = COLOR_DARK_GRAY

        # Bullet points with metrics bolded
        for bullet in job.get('bullets', []):
            clean_text, bold_parts = extract_metrics(bullet)
            add_bullet_point(doc, clean_text, bold_parts if bold_parts else None)

    # ===== EDUCATION =====
    add_horizontal_line(doc)
    add_section_header(doc, 'EDUCATION')

    for edu in education:
        # Line 1: Degree name (bold, left) — Dates (right-aligned)
        edu_para = doc.add_paragraph()
        edu_para.paragraph_format.space_before = Pt(8)
        edu_para.paragraph_format.space_after = Pt(1)
        if edu.get('dates'):
            edu_para.paragraph_format.tab_stops.add_tab_stop(
                Inches(6.3), WD_TAB_ALIGNMENT.RIGHT
            )
        degree_run = edu_para.add_run(edu['degree'])
        set_font(degree_run, size=FONT_SIZE_BODY, bold=True)
        if edu.get('dates'):
            edu_para.add_run('\t')
            dates_run = edu_para.add_run(edu['dates'])
            set_font(dates_run, size=FONT_SIZE_BODY)
            dates_run.font.color.rgb = COLOR_DARK_GRAY

        # Line 2: School, Location
        school_para = doc.add_paragraph()
        school_para.paragraph_format.space_before = Pt(0)
        school_para.paragraph_format.space_after = Pt(4)
        school_text = edu.get('school', '')
        if edu.get('location'):
            school_text += f", {edu['location']}"
        school_run = school_para.add_run(school_text)
        set_font(school_run, size=FONT_SIZE_BODY)
        school_run.font.color.rgb = COLOR_DARK_GRAY

    # ===== CERTIFICATIONS & LICENSURE =====
    if certifications:
        add_horizontal_line(doc)
        add_section_header(doc, 'CERTIFICATIONS & LICENSURE')

        # Render all certifications as a single comma-separated line
        cert_line = ', '.join(certifications)
        add_bullet_point(doc, cert_line)

    # ===== PUBLICATIONS (Optional) =====
    if publications:
        add_horizontal_line(doc)
        add_section_header(doc, 'PUBLICATIONS')

        if isinstance(publications, dict):
            # Dict format: {'Category Name': ['pub1', 'pub2'], ...}
            for category, pubs in publications.items():
                cat_para = doc.add_paragraph()
                cat_para.paragraph_format.space_before = Pt(8)
                cat_para.paragraph_format.space_after = Pt(3)
                cat_run = cat_para.add_run(category)
                set_font(cat_run, size=FONT_SIZE_BODY, bold=True)
                cat_run.font.italic = True
                for i, pub in enumerate(pubs, 1):
                    add_publication_entry(doc, pub, number=i)
        else:
            # List format: ['pub1', 'pub2', ...]
            for i, pub in enumerate(publications, 1):
                add_publication_entry(doc, pub, number=i)

        # Add publications footer (e.g., Google Scholar link) if provided
        if publications_footer:
            footer_para = doc.add_paragraph()
            footer_para.paragraph_format.space_before = Pt(6)
            footer_para.paragraph_format.space_after = Pt(2)
            footer_run = footer_para.add_run(publications_footer)
            set_font(footer_run, size=Pt(9))
            footer_run.font.color.rgb = COLOR_DARK_GRAY

    # ===== PROJECTS (Optional) =====
    if projects:
        add_horizontal_line(doc)
        add_section_header(doc, 'PROJECTS')

        for proj in projects:
            # Project title line with dates
            proj_para = doc.add_paragraph()
            proj_para.paragraph_format.space_before = Pt(8)
            proj_para.paragraph_format.space_after = Pt(2)
            title_run = proj_para.add_run(proj.get('title', ''))
            set_font(title_run, size=FONT_SIZE_BODY, bold=True)
            if proj.get('dates'):
                sep_run = proj_para.add_run('  |  ')
                set_font(sep_run, size=FONT_SIZE_BODY)
                date_run = proj_para.add_run(proj['dates'])
                set_font(date_run, size=FONT_SIZE_BODY)

            # Project bullets
            for bullet in proj.get('bullets', []):
                add_bullet_point(doc, bullet)

    # ===== PROFESSIONAL MEMBERSHIPS (Optional) =====
    if professional_memberships:
        add_horizontal_line(doc)
        add_section_header(doc, 'PROFESSIONAL MEMBERSHIPS')

        for membership in professional_memberships:
            add_bullet_point(doc, membership)

    # Save the document
    doc.save(output_path)
    return output_path


def create_ats_cover_letter(
    output_path,
    name,
    contact_info,
    date,
    recipient_info,  # dict with name, title, company, address
    job_title,
    paragraphs,  # list of paragraph strings
    closing="Sincerely,"
):
    """
    Create an ATS-optimized cover letter DOCX (Times New Roman 11pt, one page).

    Args:
        output_path: Path to save the DOCX file
        name: Full name with credentials
        contact_info: dict with city, state, phone, email
        date: Date string
        recipient_info: dict with name (optional), title (optional), company, address
        job_title: Position being applied for
        paragraphs: List of paragraph text strings
        closing: Closing phrase (default: "Sincerely,")
    """
    doc = Document()

    # Set margins (compact for one-page fit)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # ===== HEADER - Your Info =====
    header_para = doc.add_paragraph()
    header_para.paragraph_format.space_after = Pt(0)
    header_run = header_para.add_run(name)
    set_font(header_run, size=14, bold=True)
    header_run.font.color.rgb = COLOR_NAVY

    contact_para = doc.add_paragraph()
    contact_para.paragraph_format.space_before = Pt(0)
    contact_para.paragraph_format.space_after = Pt(0)
    contact_text = f"{contact_info.get('city', '')}, {contact_info.get('state', '')} | {contact_info.get('email', '')} | {contact_info.get('phone', '')}"
    contact_run = contact_para.add_run(contact_text)
    set_font(contact_run, size=FONT_SIZE_CONTACT)
    contact_run.font.color.rgb = COLOR_DARK_GRAY

    # ===== DATE =====
    date_para = doc.add_paragraph()
    date_para.paragraph_format.space_before = Pt(12)
    date_para.paragraph_format.space_after = Pt(6)
    date_run = date_para.add_run(date)
    set_font(date_run, size=FONT_SIZE_BODY)

    # ===== RECIPIENT INFO =====
    if recipient_info.get('name'):
        recip_name = doc.add_paragraph()
        recip_name_run = recip_name.add_run(recipient_info['name'])
        set_font(recip_name_run, size=FONT_SIZE_BODY)
        recip_name.paragraph_format.space_after = Pt(0)

    if recipient_info.get('title'):
        recip_title = doc.add_paragraph()
        recip_title.paragraph_format.space_before = Pt(0)
        recip_title.paragraph_format.space_after = Pt(0)
        recip_title_run = recip_title.add_run(recipient_info['title'])
        set_font(recip_title_run, size=FONT_SIZE_BODY)

    recip_company = doc.add_paragraph()
    recip_company.paragraph_format.space_before = Pt(0)
    recip_company.paragraph_format.space_after = Pt(0)
    recip_company_run = recip_company.add_run(recipient_info['company'])
    set_font(recip_company_run, size=FONT_SIZE_BODY)

    if recipient_info.get('address'):
        recip_addr = doc.add_paragraph()
        recip_addr.paragraph_format.space_before = Pt(0)
        recip_addr.paragraph_format.space_after = Pt(0)
        recip_addr_run = recip_addr.add_run(recipient_info['address'])
        set_font(recip_addr_run, size=FONT_SIZE_BODY)

    # ===== SUBJECT LINE =====
    subject_para = doc.add_paragraph()
    subject_para.paragraph_format.space_before = Pt(12)
    subject_para.paragraph_format.space_after = Pt(6)
    subject_run = subject_para.add_run(f"Re: {job_title}")
    set_font(subject_run, size=FONT_SIZE_BODY, bold=True)
    subject_run.font.color.rgb = COLOR_NAVY

    # ===== SALUTATION =====
    salutation = doc.add_paragraph()
    salutation.paragraph_format.space_after = Pt(6)
    if recipient_info.get('name'):
        sal_text = f"Dear {recipient_info['name']},"
    else:
        sal_text = "Dear Hiring Manager,"
    sal_run = salutation.add_run(sal_text)
    set_font(sal_run, size=FONT_SIZE_BODY)

    # ===== BODY PARAGRAPHS =====
    for para_text in paragraphs:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(8)

        # Bold metrics within paragraphs
        clean_text, bold_parts = extract_metrics(para_text)
        _set_line_spacing(para)

        if bold_parts:
            remaining = clean_text
            for bold_text in bold_parts:
                if bold_text in remaining:
                    parts = remaining.split(bold_text, 1)
                    if parts[0]:
                        run = para.add_run(parts[0])
                        set_font(run, size=FONT_SIZE_BODY)
                    run = para.add_run(bold_text)
                    set_font(run, size=FONT_SIZE_BODY, bold=True)
                    remaining = parts[1] if len(parts) > 1 else ''
            if remaining:
                run = para.add_run(remaining)
                set_font(run, size=FONT_SIZE_BODY)
        else:
            run = para.add_run(clean_text)
            set_font(run, size=FONT_SIZE_BODY)

    # ===== CLOSING =====
    closing_para = doc.add_paragraph()
    closing_para.paragraph_format.space_before = Pt(6)
    closing_para.paragraph_format.space_after = Pt(0)
    closing_run = closing_para.add_run(closing)
    set_font(closing_run, size=FONT_SIZE_BODY)

    # ===== SIGNATURE =====
    sig_para = doc.add_paragraph()
    sig_para.paragraph_format.space_before = Pt(24)
    sig_run = sig_para.add_run(name)
    set_font(sig_run, size=FONT_SIZE_BODY, bold=True)
    sig_run.font.color.rgb = COLOR_NAVY

    # Save
    doc.save(output_path)
    return output_path


# =============================================================================
# HARVARD STYLE RESUME GENERATOR
# =============================================================================

def add_harvard_section_header(doc, text):
    """
    Add a Harvard-style section header (Bold, underlined, not all-caps).
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, font_name='Times New Roman', size=11, bold=True)
    run.underline = True


def add_harvard_entry_with_date(doc, left_text, right_text, left_bold=False, font_size=11):
    """
    Add a line with left-aligned text and right-aligned date (using tab stop).
    This is the classic Harvard format where dates appear on the right.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Set tab stop at right margin for date alignment
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

    # Left text
    left_run = p.add_run(left_text)
    set_font(left_run, font_name='Times New Roman', size=font_size, bold=left_bold)

    # Tab and right text
    p.add_run('\t')
    right_run = p.add_run(right_text)
    set_font(right_run, font_name='Times New Roman', size=font_size)

    return p


def add_harvard_bullet(doc, text, bold_parts=None, font_size=11):
    """
    Add a Harvard-style bullet point with Times New Roman font.
    """
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25)

    if bold_parts:
        remaining = text
        for bold_text in bold_parts:
            if bold_text in remaining:
                parts = remaining.split(bold_text, 1)
                if parts[0]:
                    run = p.add_run(parts[0])
                    set_font(run, font_name='Times New Roman', size=font_size)
                run = p.add_run(bold_text)
                set_font(run, font_name='Times New Roman', size=font_size, bold=True)
                remaining = parts[1] if len(parts) > 1 else ''
        if remaining:
            run = p.add_run(remaining)
            set_font(run, font_name='Times New Roman', size=font_size)
    else:
        run = p.add_run(text)
        set_font(run, font_name='Times New Roman', size=font_size)

    return p


def create_harvard_resume(
    output_path,
    name,
    contact_info,  # dict with address, city, state, zip, phone, email, linkedin
    education,  # list of dicts - comes FIRST in Harvard format
    experience,  # list of dicts with title, company, location, dates, bullets
    skills=None,  # dict with categories like 'languages', 'technical', 'interests'
    certifications=None,  # list of strings
    publications=None,  # list of strings
    leadership=None,  # list of dicts for leadership/activities
    honors=None,  # list of strings for awards/honors
    include_summary=False,  # Harvard typically doesn't include summary for students
    summary=None,
    core_competencies=None  # list of keywords - appears after summary if provided
):
    """
    Create a Harvard OCS-style resume DOCX.

    Harvard Format Key Features:
    - Times New Roman font throughout
    - Education section comes FIRST
    - Dates are right-aligned
    - Company/Organization names are bold
    - Clean, traditional academic look
    - No graphics or fancy formatting

    Args:
        output_path: Path to save the DOCX file
        name: Full name with credentials
        contact_info: dict with address, city, state, zip, phone, email, linkedin
        education: List of education dicts with: school, degree, location, dates, gpa, honors, coursework
        experience: List of job dicts with: title, company, location, dates, bullets
        skills: Optional dict with skill categories
        certifications: Optional list of certification strings
        publications: Optional list of publication strings
        leadership: Optional list of leadership/activity dicts
        honors: Optional list of award/honor strings
        include_summary: Whether to include a summary (False for traditional Harvard)
        summary: Optional summary text
    """
    doc = Document()

    # Set margins (Harvard uses 0.5" to 1")
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # ===== NAME (Centered, Bold, 14pt) =====
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(0)
    name_run = name_para.add_run(name)
    set_font(name_run, font_name='Times New Roman', size=14, bold=True)

    # ===== CONTACT INFO (Centered, single or two lines) =====
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_before = Pt(0)
    contact_para.paragraph_format.space_after = Pt(0)

    # Build contact line
    contact_parts = []
    if contact_info.get('address'):
        contact_parts.append(contact_info['address'])
    if contact_info.get('city') and contact_info.get('state'):
        city_state = f"{contact_info['city']}, {contact_info['state']}"
        if contact_info.get('zip'):
            city_state += f" {contact_info['zip']}"
        contact_parts.append(city_state)

    if contact_parts:
        address_run = contact_para.add_run(' | '.join(contact_parts))
        set_font(address_run, font_name='Times New Roman', size=10)

    # Second line: phone, email, linkedin
    contact_para2 = doc.add_paragraph()
    contact_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para2.paragraph_format.space_before = Pt(0)
    contact_para2.paragraph_format.space_after = Pt(6)

    contact_line2_parts = []
    if contact_info.get('phone'):
        contact_line2_parts.append(contact_info['phone'])
    if contact_info.get('email'):
        contact_line2_parts.append(contact_info['email'])
    if contact_info.get('linkedin'):
        contact_line2_parts.append(contact_info['linkedin'])

    contact_run2 = contact_para2.add_run(' | '.join(contact_line2_parts))
    set_font(contact_run2, font_name='Times New Roman', size=10)

    # ===== SUMMARY (Optional - not traditional Harvard) =====
    if include_summary and summary:
        add_harvard_section_header(doc, 'Summary')
        summary_para = doc.add_paragraph()
        summary_para.paragraph_format.space_after = Pt(6)
        summary_run = summary_para.add_run(summary)
        set_font(summary_run, font_name='Times New Roman', size=11)

    # ===== CORE COMPETENCIES (After summary, before education) =====
    if core_competencies:
        add_harvard_section_header(doc, 'Core Competencies')

        # Display as bullet list
        for competency in core_competencies:
            add_harvard_bullet(doc, competency)

    # ===== EDUCATION (Comes FIRST in Harvard format) =====
    add_harvard_section_header(doc, 'Education')

    for edu in education:
        # School name (bold) with dates right-aligned
        school_line = add_harvard_entry_with_date(
            doc,
            edu.get('school', ''),
            edu.get('dates', ''),
            left_bold=True
        )

        # Degree and location on second line
        if edu.get('degree'):
            degree_para = doc.add_paragraph()
            degree_para.paragraph_format.space_before = Pt(0)
            degree_para.paragraph_format.space_after = Pt(0)

            degree_run = degree_para.add_run(edu['degree'])
            set_font(degree_run, font_name='Times New Roman', size=11)

            if edu.get('location'):
                loc_run = degree_para.add_run(f", {edu['location']}")
                set_font(loc_run, font_name='Times New Roman', size=11)

        # GPA if provided
        if edu.get('gpa'):
            gpa_para = doc.add_paragraph()
            gpa_para.paragraph_format.space_before = Pt(0)
            gpa_para.paragraph_format.space_after = Pt(0)
            gpa_run = gpa_para.add_run(f"GPA: {edu['gpa']}")
            set_font(gpa_run, font_name='Times New Roman', size=11)

        # Honors if provided
        if edu.get('honors'):
            honors_para = doc.add_paragraph()
            honors_para.paragraph_format.space_before = Pt(0)
            honors_para.paragraph_format.space_after = Pt(0)
            if isinstance(edu['honors'], list):
                honors_text = ', '.join(edu['honors'])
            else:
                honors_text = edu['honors']
            honors_run = honors_para.add_run(f"Honors: {honors_text}")
            set_font(honors_run, font_name='Times New Roman', size=11)

        # Relevant coursework if provided
        if edu.get('coursework'):
            course_para = doc.add_paragraph()
            course_para.paragraph_format.space_before = Pt(0)
            course_para.paragraph_format.space_after = Pt(4)
            if isinstance(edu['coursework'], list):
                course_text = ', '.join(edu['coursework'])
            else:
                course_text = edu['coursework']
            course_run = course_para.add_run(f"Relevant Coursework: {course_text}")
            set_font(course_run, font_name='Times New Roman', size=10)

    # ===== EXPERIENCE =====
    add_harvard_section_header(doc, 'Experience')

    for job in experience:
        # Company/Organization name (bold) with dates right-aligned
        add_harvard_entry_with_date(
            doc,
            job.get('company', ''),
            job.get('dates', ''),
            left_bold=True
        )

        # Job title and location on second line (not bold)
        title_para = doc.add_paragraph()
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after = Pt(2)

        title_run = title_para.add_run(job.get('title', ''))
        set_font(title_run, font_name='Times New Roman', size=11)
        title_run.italic = True  # Harvard often italicizes titles

        if job.get('location'):
            loc_run = title_para.add_run(f", {job['location']}")
            set_font(loc_run, font_name='Times New Roman', size=11)

        # Bullet points
        for bullet in job.get('bullets', []):
            clean_text, bold_parts = extract_metrics(bullet)
            add_harvard_bullet(doc, clean_text, bold_parts if bold_parts else None)

    # ===== LEADERSHIP & ACTIVITIES (Optional) =====
    if leadership:
        add_harvard_section_header(doc, 'Leadership & Activities')

        for activity in leadership:
            add_harvard_entry_with_date(
                doc,
                activity.get('organization', ''),
                activity.get('dates', ''),
                left_bold=True
            )

            if activity.get('role'):
                role_para = doc.add_paragraph()
                role_para.paragraph_format.space_before = Pt(0)
                role_para.paragraph_format.space_after = Pt(2)
                role_run = role_para.add_run(activity['role'])
                set_font(role_run, font_name='Times New Roman', size=11)
                role_run.italic = True

            for bullet in activity.get('bullets', []):
                add_harvard_bullet(doc, bullet)

    # ===== PUBLICATIONS (Optional) =====
    if publications:
        add_harvard_section_header(doc, 'Publications')
        for pub in publications:
            add_harvard_bullet(doc, pub, font_size=10)

    # ===== HONORS & AWARDS (Optional) =====
    if honors:
        add_harvard_section_header(doc, 'Honors & Awards')
        for honor in honors:
            add_harvard_bullet(doc, honor)

    # ===== CERTIFICATIONS (Optional) =====
    if certifications:
        add_harvard_section_header(doc, 'Certifications')
        for cert in certifications:
            add_harvard_bullet(doc, cert)

    # ===== SKILLS & ADDITIONAL (Optional) =====
    if skills:
        add_harvard_section_header(doc, 'Skills & Additional')

        for category, items in skills.items():
            skill_para = doc.add_paragraph()
            skill_para.paragraph_format.space_before = Pt(0)
            skill_para.paragraph_format.space_after = Pt(2)

            # Category name (bold)
            cat_run = skill_para.add_run(f"{category}: ")
            set_font(cat_run, font_name='Times New Roman', size=11, bold=True)

            # Skills list
            if isinstance(items, list):
                items_text = ', '.join(items)
            else:
                items_text = items
            items_run = skill_para.add_run(items_text)
            set_font(items_run, font_name='Times New Roman', size=11)

    # Save the document
    doc.save(output_path)
    return output_path


def create_harvard_cover_letter(
    output_path,
    name,
    contact_info,
    date,
    recipient_info,
    job_title,
    paragraphs,
    closing="Sincerely,"
):
    """
    Create a Harvard-style cover letter DOCX.

    Args:
        output_path: Path to save the DOCX file
        name: Full name with credentials
        contact_info: dict with city, state, phone, email
        date: Date string
        recipient_info: dict with name, title, company, address
        job_title: Position being applied for
        paragraphs: List of paragraph text strings
        closing: Closing phrase (default: "Sincerely,")
    """
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ===== YOUR CONTACT INFO (Top right or centered) =====
    header_para = doc.add_paragraph()
    header_run = header_para.add_run(name)
    set_font(header_run, font_name='Times New Roman', size=12, bold=True)

    # Address
    if contact_info.get('address') or contact_info.get('city'):
        addr_para = doc.add_paragraph()
        addr_para.paragraph_format.space_before = Pt(0)
        addr_para.paragraph_format.space_after = Pt(0)
        addr_text = ""
        if contact_info.get('address'):
            addr_text = contact_info['address']
        if contact_info.get('city'):
            if addr_text:
                addr_text += ", "
            addr_text += f"{contact_info.get('city', '')}, {contact_info.get('state', '')}"
            if contact_info.get('zip'):
                addr_text += f" {contact_info['zip']}"
        addr_run = addr_para.add_run(addr_text)
        set_font(addr_run, font_name='Times New Roman', size=11)

    # Phone and email
    contact_para = doc.add_paragraph()
    contact_para.paragraph_format.space_before = Pt(0)
    contact_parts = []
    if contact_info.get('phone'):
        contact_parts.append(contact_info['phone'])
    if contact_info.get('email'):
        contact_parts.append(contact_info['email'])
    contact_run = contact_para.add_run(' | '.join(contact_parts))
    set_font(contact_run, font_name='Times New Roman', size=11)

    # ===== DATE =====
    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(date)
    set_font(date_run, font_name='Times New Roman', size=11)

    # ===== RECIPIENT INFO =====
    doc.add_paragraph()

    if recipient_info.get('name'):
        recip_name = doc.add_paragraph()
        recip_name.paragraph_format.space_after = Pt(0)
        recip_name_run = recip_name.add_run(recipient_info['name'])
        set_font(recip_name_run, font_name='Times New Roman', size=11)

    if recipient_info.get('title'):
        recip_title = doc.add_paragraph()
        recip_title.paragraph_format.space_before = Pt(0)
        recip_title.paragraph_format.space_after = Pt(0)
        recip_title_run = recip_title.add_run(recipient_info['title'])
        set_font(recip_title_run, font_name='Times New Roman', size=11)

    recip_company = doc.add_paragraph()
    recip_company.paragraph_format.space_before = Pt(0)
    recip_company.paragraph_format.space_after = Pt(0)
    recip_company_run = recip_company.add_run(recipient_info.get('company', ''))
    set_font(recip_company_run, font_name='Times New Roman', size=11)

    if recipient_info.get('address'):
        recip_addr = doc.add_paragraph()
        recip_addr.paragraph_format.space_before = Pt(0)
        recip_addr_run = recip_addr.add_run(recipient_info['address'])
        set_font(recip_addr_run, font_name='Times New Roman', size=11)

    # ===== SALUTATION =====
    doc.add_paragraph()
    salutation = doc.add_paragraph()
    if recipient_info.get('name'):
        sal_text = f"Dear {recipient_info['name']},"
    else:
        sal_text = "Dear Hiring Manager,"
    sal_run = salutation.add_run(sal_text)
    set_font(sal_run, font_name='Times New Roman', size=11)

    # ===== BODY PARAGRAPHS =====
    for para_text in paragraphs:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.first_line_indent = Inches(0.5)  # Harvard uses indented paragraphs

        run = para.add_run(para_text)
        set_font(run, font_name='Times New Roman', size=11)

    # ===== CLOSING =====
    doc.add_paragraph()
    closing_para = doc.add_paragraph()
    closing_run = closing_para.add_run(closing)
    set_font(closing_run, font_name='Times New Roman', size=11)

    # ===== SIGNATURE =====
    doc.add_paragraph()
    doc.add_paragraph()
    sig_para = doc.add_paragraph()
    sig_run = sig_para.add_run(name)
    set_font(sig_run, font_name='Times New Roman', size=11)

    # Save
    doc.save(output_path)
    return output_path


# ===== EXAMPLE USAGE =====
if __name__ == '__main__':
    # Example data structure for testing
    example_resume = {
        'name': 'JANE DOE, M.D.',
        'contact_info': {
            'city': 'New York',
            'state': 'NY',
            'zip': '10001',
            'phone': '555-123-4567',
            'email': 'jane.doe@example.com',
            'linkedin': 'linkedin.com/in/janedoe'
        },
        'summary': 'Results-oriented Physician and Clinical Research Professional with over 10 years of experience in Clinical Development and Drug Safety. Proven track record in Signal Detection, Risk Management, and Cross-functional Team Leadership. Committed to driving patient safety and regulatory compliance in global pharmaceutical development.',
        'core_competencies': [
            'Clinical Research',
            'Signal Detection & Analysis',
            'Risk Management & Benefit-Risk Assessment',
            'Safety Data Review',
            'FDA/ICH-GCP Compliance',
            'Cross-functional Team Leadership',
            'Medical Monitoring',
            'Regulatory Reporting',
            'Health Authority Interactions',
            'Real World Evidence (RWE)'
        ],
        'experience': [
            {
                'title': 'Graduate Researcher',
                'company': 'University Medical Center',
                'location': 'New York, NY',
                'dates': 'June 2024 – Present',
                'bullets': [
                    'Designed and deployed a deep learning model using Real World Data (11,300+ records) to predict clinical outcomes with 91% sensitivity.',
                    'Engineered utility-driven alert threshold strategy to optimize sensitivity-specificity trade-off for safety signal detection.',
                    'Authored and submitted IRB protocol for human subjects research, defining data management plans and safety escalation protocols.'
                ]
            },
            {
                'title': 'Clinical Operations Lead',
                'company': 'Regional Medical Center',
                'location': 'Newark, NJ',
                'dates': 'July 2023 – May 2024',
                'bullets': [
                    'Directed clinical site operations across 8 centers, ensuring 100% adherence to FDA/GCP safety protocols.',
                    'Led aggregate safety data review and assessed AE/SAEs in partnership with Principal Investigators.',
                    'Standardized scientific communication workflows to accelerate data cleaning and interim safety analysis.'
                ]
            }
        ],
        'education': [
            {
                'degree': 'Master of Public Health, Health Informatics',
                'school': 'University Medical School',
                'location': 'New York, NY',
                'dates': '2024 – 2026'
            },
            {
                'degree': 'Bachelor of Medicine, Bachelor of Surgery (M.B.B.S.)',
                'school': 'School of Medicine',
                'location': 'International',
                'dates': '2007 – 2014'
            }
        ],
        'certifications': [
            'ACLS & BLS Certified',
            'CITI Human Subjects Research (HSR)',
            'Clinical Trials Operations Certification',
            'Board Certified'
        ],
        'professional_memberships': [
            'American Medical Association',
            'American Public Health Association'
        ]
    }


# =============================================================================
# MARKDOWN-TO-DOCX PARSERS
# =============================================================================

def _parse_school_location(school_loc: str, edu_entry: dict):
    """
    Parse a school/location string into edu_entry['school'] and edu_entry['location'].

    Handles formats like:
      "Yale School of Public Health, New Haven, CT"  -> school="Yale...", location="New Haven, CT"
      "Faculty of Medicine, University of Jaffna, Sri Lanka" -> school="...", location="Sri Lanka"
      "Harvard University" -> school="Harvard University", location=""

    The naive rsplit(',', 1) approach incorrectly splits city names into the school field
    when the pattern is "School, City, ST". This function detects the US state-code or
    country at the end and splits accordingly so Workday can recognize the institution name.
    """
    school_loc = school_loc.strip()
    if not school_loc:
        return

    # Pattern 1: ends with ", XX" where XX is a 2-letter US state code (all uppercase)
    # e.g. "Yale School of Public Health, New Haven, CT"
    state_match = re.search(r',\s*([A-Z]{2})\s*$', school_loc)
    if state_match:
        # Everything before the state = "School, City"
        school_and_city = school_loc[:state_match.start()].strip()
        state = state_match.group(1)
        # Split on last comma to separate city from school name
        last_comma = school_and_city.rfind(',')
        if last_comma != -1:
            edu_entry['school'] = school_and_city[:last_comma].strip()
            city = school_and_city[last_comma + 1:].strip()
            edu_entry['location'] = f"{city}, {state}"
        else:
            edu_entry['school'] = school_and_city
            edu_entry['location'] = state
        return

    # Pattern 2: no state code — use rsplit on last comma (handles "School, Country")
    if ',' in school_loc:
        school_parts = school_loc.rsplit(',', 1)
        edu_entry['school'] = school_parts[0].strip()
        edu_entry['location'] = school_parts[1].strip()
    else:
        edu_entry['school'] = school_loc


_KNOWN_SECTION_HEADERS = [
    'PROFESSIONAL SUMMARY', 'SUMMARY', 'CORE COMPETENCIES', 'COMPETENCIES',
    'PROFESSIONAL EXPERIENCE', 'WORK EXPERIENCE', 'EXPERIENCE',
    'EDUCATION', 'CERTIFICATIONS', 'LICENSURE', 'PUBLICATIONS',
    'PROJECTS', 'PROFESSIONAL MEMBERSHIPS', 'MEMBERSHIPS', 'AWARDS', 'HONORS',
    'SKILLS', 'LANGUAGES', 'VOLUNTEER',
]


def _is_known_section_header(line: str) -> bool:
    upper = line.strip().upper()
    return any(h in upper for h in _KNOWN_SECTION_HEADERS)


def _preprocess_resume_md(md_text: str) -> str:
    """Remove ___ separator lines that are NOT immediately before a known section header.

    Resumes often use ___ between individual job entries within PROFESSIONAL EXPERIENCE.
    Those separators cause parse_resume_markdown to treat each job as its own section,
    silently dropping all but the first job. This function removes intra-section separators
    while preserving inter-section separators.
    """
    lines = md_text.split('\n')
    result = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if re.match(r'^_{3,}\s*$', line.strip()):
            # Look ahead past blank lines to find the next non-blank line
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and _is_known_section_header(lines[j]):
                result.append(line)  # keep: separates two real sections
            # else: skip (intra-section separator between jobs)
        else:
            result.append(line)
        i += 1
    return '\n'.join(result)


def parse_resume_markdown(md_text: str) -> dict:
    """
    Parse ATS/Workday-format resume markdown into a dict matching create_ats_resume() params.

    Returns dict with keys: name, contact_info, summary, core_competencies,
                            experience, education, certifications,
                            professional_memberships, publications
    """
    result = {
        'name': '',
        'contact_info': {},
        'summary': '',
        'core_competencies': [],
        'experience': [],
        'education': [],
        'certifications': [],
        'professional_memberships': [],
        'publications': None,
        'projects': None,
    }

    # Remove intra-section separators (between jobs) before splitting
    md_text = _preprocess_resume_md(md_text)

    # Split into sections by separator lines (3+ underscores)
    sections = re.split(r'\n_{3,}\n', md_text)

    if not sections:
        return result

    # -- Section 0: Header (name, contact, linkedin) --
    header_lines = [l.strip() for l in sections[0].strip().split('\n') if l.strip()]
    if header_lines:
        result['name'] = header_lines[0]

    if len(header_lines) >= 2:
        contact_line = header_lines[1]
        parts = [p.strip() for p in contact_line.split('|')]
        ci = {}
        for part in parts:
            if '@' in part:
                ci['email'] = part.strip()
            elif re.search(r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}', part):
                ci['phone'] = part.strip()
            else:
                # Location: "City, State ZIP"
                loc_match = re.match(r'(.+?),\s*([A-Z]{2})\s*(\d{5})?', part)
                if loc_match:
                    ci['city'] = loc_match.group(1).strip()
                    ci['state'] = loc_match.group(2).strip()
                    if loc_match.group(3):
                        ci['zip'] = loc_match.group(3).strip()
        result['contact_info'] = ci

    if len(header_lines) >= 3:
        for line in header_lines[2:]:
            if 'linkedin' in line.lower() or 'linkedin.com' in line.lower():
                result['contact_info']['linkedin'] = line.strip()
                break

    # -- Process remaining sections by header --
    for section in sections[1:]:
        lines = section.strip().split('\n')
        if not lines:
            continue

        # First non-empty line is the section header
        header = lines[0].strip().upper()
        body_lines = lines[1:]

        if 'PROFESSIONAL SUMMARY' in header or 'SUMMARY' in header:
            result['summary'] = '\n'.join(l.strip() for l in body_lines if l.strip())

        elif 'CORE COMPETENCIES' in header or 'COMPETENCIES' in header:
            comps = []
            for line in body_lines:
                line = line.strip()
                if not line:
                    continue
                # Split on bullet char or multi-space gaps
                if '\u2022' in line:
                    parts = [p.strip() for p in line.split('\u2022') if p.strip()]
                    comps.extend(parts)
                elif '- ' in line:
                    parts = [p.strip().lstrip('- ') for p in re.split(r'\s{2,}', line) if p.strip()]
                    comps.extend(parts)
                else:
                    # Try splitting on multiple spaces (multi-column layout)
                    parts = [p.strip() for p in re.split(r'\s{3,}', line) if p.strip()]
                    if len(parts) > 1:
                        comps.extend(parts)
                    elif line:
                        comps.append(line)
            result['core_competencies'] = comps

        elif 'PROFESSIONAL EXPERIENCE' in header or 'EXPERIENCE' in header:
            jobs = []
            current_job = None
            pending_title = None  # for TITLE\nCOMPANY | LOCATION two-line format
            _date_re = re.compile(
                r'(January|February|March|April|May|June|July|August|September|'
                r'October|November|December|Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\d{4})'
            )
            for line in body_lines:
                stripped = line.strip()
                if not stripped:
                    continue

                pipe_parts = [p.strip() for p in stripped.split('|')]
                is_bullet = stripped.startswith('\u2022') or stripped.startswith('-')
                has_date = bool(_date_re.search(stripped))

                if len(pipe_parts) >= 3 and not is_bullet:
                    # TITLE | COMPANY | Location (single-line format)
                    if current_job:
                        jobs.append(current_job)
                    current_job = {
                        'title': pipe_parts[0],
                        'company': pipe_parts[1],
                        'location': pipe_parts[2],
                        'dates': '',
                        'bullets': [],
                    }
                    pending_title = None
                elif len(pipe_parts) == 2 and not is_bullet and pending_title is not None:
                    # COMPANY | Location following a standalone title line (two-line format)
                    if current_job:
                        jobs.append(current_job)
                    current_job = {
                        'title': pending_title,
                        'company': pipe_parts[0],
                        'location': pipe_parts[1],
                        'dates': '',
                        'bullets': [],
                    }
                    pending_title = None
                elif not is_bullet and has_date:
                    # Date line
                    if current_job and not current_job['dates']:
                        current_job['dates'] = stripped
                    pending_title = None
                elif is_bullet:
                    if current_job:
                        bullet = re.sub(r'^[\u2022\-]\s*', '', stripped)
                        current_job['bullets'].append(bullet)
                    pending_title = None
                elif len(pipe_parts) == 1 and not is_bullet and not has_date:
                    # No pipes, not a bullet, not a date — potential standalone title
                    pending_title = stripped
                elif current_job:
                    # Continuation line inside a job
                    if len(stripped) > 10 and not stripped.isupper():
                        current_job['bullets'].append(stripped)
                    pending_title = None

            if current_job:
                jobs.append(current_job)
            result['experience'] = jobs

        elif 'EDUCATION' in header:
            edu_list = []
            i = 0
            edu_lines = [l.strip() for l in body_lines if l.strip()]
            while i < len(edu_lines):
                line = edu_lines[i]
                # Degree line (doesn't start with bullet, not a date)
                if not line.startswith('\u2022') and not line.startswith('-'):
                    edu_entry = {'degree': line, 'school': '', 'location': '', 'dates': ''}
                    if i + 1 < len(edu_lines):
                        next_line = edu_lines[i + 1]
                        # School line may contain | for dates or location
                        if '|' in next_line:
                            parts = [p.strip() for p in next_line.split('|')]
                            school_loc = parts[0]
                            if len(parts) >= 2:
                                edu_entry['dates'] = parts[-1].strip()
                            # Parse "School, City, ST" or "School, Country" or just "School"
                            _parse_school_location(school_loc, edu_entry)
                        else:
                            # Check if next line has comma (School, Location)
                            _parse_school_location(next_line, edu_entry)
                        i += 2
                    else:
                        i += 1

                    # Check for GPA or additional info lines
                    while i < len(edu_lines) and (edu_lines[i].startswith('\u2022') or edu_lines[i].startswith('-') or 'GPA' in edu_lines[i].upper()):
                        i += 1

                    edu_list.append(edu_entry)
                else:
                    i += 1
            result['education'] = edu_list

        elif 'CERTIFICATIONS' in header or 'LICENSURE' in header:
            certs = []
            for line in body_lines:
                stripped = line.strip()
                if stripped and (stripped.startswith('\u2022') or stripped.startswith('-')):
                    certs.append(re.sub(r'^[\u2022\-]\s*', '', stripped))
                elif stripped and len(stripped) > 3:
                    certs.append(stripped)
            result['certifications'] = certs

        elif 'PUBLICATIONS' in header:
            pubs = {}
            current_category = None
            pub_footer = None
            for line in body_lines:
                stripped = line.strip()
                if not stripped:
                    continue
                # Detect "Full publication list:" or similar footer with URL
                if stripped.lower().startswith('full publication') or stripped.lower().startswith('full list'):
                    pub_footer = stripped
                    continue
                if stripped.startswith('\u2022') or stripped.startswith('-'):
                    pub = re.sub(r'^[\u2022\-]\s*', '', stripped)
                    if current_category:
                        if current_category not in pubs:
                            pubs[current_category] = []
                        pubs[current_category].append(pub)
                    else:
                        if '_default' not in pubs:
                            pubs['_default'] = []
                        pubs['_default'].append(pub)
                else:
                    # Non-bullet line = subcategory header
                    if not stripped.isupper():
                        current_category = stripped
            # Extract footer before post-processing categories
            footer = pub_footer
            # Count real categories (not _default, not _footer)
            real_cats = {k: v for k, v in pubs.items() if k not in ('_default', '_footer')}

            if '_default' in pubs and not real_cats:
                # All pubs are flat (no sub-headers) — return as list
                result['publications'] = pubs['_default']
            elif pubs:
                if '_default' in pubs:
                    first_cat = next((k for k in pubs if k not in ('_default', '_footer')), None)
                    if first_cat:
                        pubs[first_cat] = pubs.pop('_default') + pubs[first_cat]
                    else:
                        result['publications'] = pubs.pop('_default')
                # Remove internal keys before passing as dict
                pubs.pop('_footer', None)
                if pubs:
                    result['publications'] = pubs
            if footer:
                result['publications_footer'] = footer

        elif 'PROJECTS' in header and 'EXPERIENCE' not in header:
            projects = []
            current_project = None
            for line in body_lines:
                stripped = line.strip()
                if not stripped:
                    continue
                is_bullet = stripped.startswith('\u2022') or stripped.startswith('-')
                if not is_bullet and '|' in stripped:
                    # Project title line: "Title | Dates"
                    if current_project:
                        projects.append(current_project)
                    parts = [p.strip() for p in stripped.split('|')]
                    current_project = {
                        'title': parts[0],
                        'dates': parts[1] if len(parts) > 1 else '',
                        'bullets': [],
                    }
                elif is_bullet and current_project:
                    bullet = re.sub(r'^[\u2022\-]\s*', '', stripped)
                    current_project['bullets'].append(bullet)
                elif not is_bullet and current_project:
                    # Continuation line
                    current_project['bullets'].append(stripped)
            if current_project:
                projects.append(current_project)
            result['projects'] = projects

        elif 'PROFESSIONAL MEMBERSHIPS' in header or 'MEMBERSHIPS' in header:
            members = []
            for line in body_lines:
                stripped = line.strip()
                if stripped and (stripped.startswith('\u2022') or stripped.startswith('-')):
                    members.append(re.sub(r'^[\u2022\-]\s*', '', stripped))
                elif stripped and len(stripped) > 3:
                    members.append(stripped)
            result['professional_memberships'] = members

    return result


def create_resume_from_md(md_path: str, output_path: str) -> str:
    """
    Read a resume markdown file and create an ATS-compliant DOCX.

    Args:
        md_path: Path to the resume .md file
        output_path: Path to save the .docx file

    Returns:
        Path to created DOCX file
    """
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    data = parse_resume_markdown(md_text)

    return create_ats_resume(
        output_path=output_path,
        name=data['name'],
        contact_info=data['contact_info'],
        summary=data['summary'],
        core_competencies=data['core_competencies'],
        experience=data['experience'],
        education=data['education'],
        certifications=data['certifications'],
        professional_memberships=data.get('professional_memberships'),
        publications=data.get('publications'),
        publications_footer=data.get('publications_footer'),
        projects=data.get('projects'),
    )


def parse_cover_letter_markdown(md_text: str) -> dict:
    """
    Parse cover letter markdown into a dict matching create_ats_cover_letter() params.

    Returns dict with keys: name, contact_info, date, recipient_info, job_title, paragraphs, closing
    """
    result = {
        'name': '',
        'contact_info': {},
        'date': '',
        'recipient_info': {},
        'job_title': '',
        'paragraphs': [],
        'closing': 'Sincerely,',
    }

    lines = md_text.strip().split('\n')

    # Find key anchor points
    dear_idx = None
    closing_idx = None
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped.lower().startswith('dear'):
            dear_idx = i
        if stripped.lower().startswith('sincerely') or stripped.lower().startswith('best regards') or stripped.lower().startswith('respectfully'):
            closing_idx = i

    if dear_idx is None:
        # Can't parse without a greeting
        result['paragraphs'] = [md_text]
        return result

    # -- Before "Dear" line: header info --
    header_lines = [l.strip() for l in lines[:dear_idx] if l.strip()]

    if header_lines:
        result['name'] = header_lines[0]

    # Parse contact info from header
    ci = {}
    for line in header_lines[1:]:
        if '@' in line and '|' in line:
            # Phone | Email line
            parts = [p.strip() for p in line.split('|')]
            for part in parts:
                if '@' in part:
                    ci['email'] = part
                elif re.search(r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}', part):
                    ci['phone'] = part
        elif '@' in line:
            ci['email'] = line.strip()
        elif re.search(r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}', line):
            ci['phone'] = line.strip()
        elif re.match(r'.*,\s*[A-Z]{2}\s*\d{5}', line):
            # Address line: City, ST ZIP
            loc_match = re.match(r'(.+?),\s*([A-Z]{2})\s*(\d{5})?', line)
            if loc_match:
                ci['city'] = loc_match.group(1).strip()
                ci['state'] = loc_match.group(2).strip()
                if loc_match.group(3):
                    ci['zip'] = loc_match.group(3).strip()
        elif re.match(r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s*\d{4}', line):
            result['date'] = line.strip()
        elif not result['date'] and re.match(r'\w+\s+\d{1,2},?\s*\d{4}', line):
            result['date'] = line.strip()
        else:
            # Could be company name or address for recipient
            if not result['recipient_info'].get('company'):
                result['recipient_info']['company'] = line.strip()
            elif not result['recipient_info'].get('address'):
                result['recipient_info']['address'] = line.strip()

    result['contact_info'] = ci

    # -- Greeting line: extract job title hint --
    greeting = lines[dear_idx].strip()
    # Try to extract job title from greeting if present
    title_match = re.search(r'(?:for the|regarding the|for your)\s+(.+?)\s+(?:position|role|opportunity)', greeting, re.IGNORECASE)
    if title_match:
        result['job_title'] = title_match.group(1).strip()

    # -- Body paragraphs (between Dear and closing) --
    if closing_idx:
        body_text = '\n'.join(lines[dear_idx + 1:closing_idx])
        result['closing'] = lines[closing_idx].strip()
    else:
        body_text = '\n'.join(lines[dear_idx + 1:])

    # Split on blank lines to get paragraphs
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', body_text) if p.strip()]
    result['paragraphs'] = paragraphs

    return result


def create_cover_letter_from_md(md_path: str, output_path: str, job_title: str = '') -> str:
    """
    Read a cover letter markdown file and create an ATS-compliant DOCX.

    Args:
        md_path: Path to the cover letter .md file
        output_path: Path to save the .docx file
        job_title: Job title (overrides parsed value if provided)

    Returns:
        Path to created DOCX file
    """
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    data = parse_cover_letter_markdown(md_text)

    # Use provided job_title if given, else fall back to parsed
    final_job_title = job_title if job_title else data.get('job_title', '')

    return create_ats_cover_letter(
        output_path=output_path,
        name=data['name'],
        contact_info=data['contact_info'],
        date=data['date'],
        recipient_info=data['recipient_info'],
        job_title=final_job_title,
        paragraphs=data['paragraphs'],
        closing=data.get('closing', 'Sincerely,'),
    )


if __name__ == '__main__':
    # Test resume generation
    create_ats_resume(
        'test_resume.docx',
        example_resume['name'],
        example_resume['contact_info'],
        example_resume['summary'],
        example_resume['core_competencies'],
        example_resume['experience'],
        example_resume['education'],
        example_resume['certifications'],
        example_resume['professional_memberships']
    )
    print("Test resume created: test_resume.docx")
