#!/usr/bin/env python3
"""
Legacy Resume Builder CLI Guard
-------------------------------
Direct resume tailoring is disabled. Use the native Resume Team so candidate
fit, provenance, audit, and publication authorization are enforced.

Usage:
    python resume_builder.py

Or with Claude Code custom commands:
    /resume
    /cover-letter
"""

from __future__ import annotations

import json
import os
import re
import sys
import tempfile
from datetime import datetime
from pathlib import Path

from legacy_rewrite_guard import (
    NATIVE_RESUME_TEAM_REQUIRED,
    NATIVE_RESUME_TEAM_REQUIRED_MESSAGE,
    NativeResumeTeamRequiredError,
)

try:
    import anthropic
except ImportError:
    anthropic = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError:
    Document = Pt = Inches = WD_ALIGN_PARAGRAPH = None


# Configuration
CONFIG_FILE = Path(__file__).parent / "config.json"
DEFAULT_CONFIG = {
    "master_resume_path": "YOUR_MASTER_RESUME.docx",
    "output_base_dir": "applications",
    "user_name": "Your Name",
    "user_credentials": "",
    "user_email": "your.email@example.com",
    "user_phone": "555-123-4567",
    "user_linkedin": "linkedin.com/in/your-profile",
    "user_city": "Your City",
    "user_state": "ST",
    "user_zip": "00000"
}

# Values that indicate the user hasn't customized their config
_PLACEHOLDER_VALUES = {
    "master_resume_path": "YOUR_MASTER_RESUME.docx",
    "user_name": "Your Name",
    "user_email": "your.email@example.com",
    "user_phone": "555-123-4567",
    "user_linkedin": "linkedin.com/in/your-profile",
    "user_city": "Your City",
    "user_state": "ST",
    "user_zip": "00000",
}


def validate_config(config):
    """Check if any config fields still have default placeholder values and warn the user."""
    warnings = []
    for field, placeholder in _PLACEHOLDER_VALUES.items():
        if config.get(field) == placeholder:
            warnings.append(f"  - {field}: still set to default '{placeholder}'")
    if warnings:
        print("WARNING: The following config.json fields have not been customized:")
        for w in warnings:
            print(w)
        print("Edit config.json to set your personal details.\n")


def load_config():
    """Load configuration from config.json or create default."""
    if CONFIG_FILE.exists():
        with open(CONFIG_FILE, 'r') as f:
            config = json.load(f)
        # Backfill any new keys from DEFAULT_CONFIG
        updated = False
        for key, value in DEFAULT_CONFIG.items():
            if key not in config:
                config[key] = value
                updated = True
        if updated:
            save_config(config)
        validate_config(config)
        return config
    else:
        save_config(DEFAULT_CONFIG)
        validate_config(DEFAULT_CONFIG)
        return DEFAULT_CONFIG


def save_config(config):
    """Save configuration to config.json."""
    with open(CONFIG_FILE, 'w') as f:
        json.dump(config, f, indent=2)


def extract_text_from_file(file_path: str) -> str:
    """Extract text content from a resume file (PDF, DOCX, MD, TXT)."""
    ext = Path(file_path).suffix.lower()
    text = ""
    try:
        if ext == ".pdf":
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        elif ext == ".docx":
            doc = Document(file_path)
            text = "\n".join(para.text for para in doc.paragraphs)
        elif ext in (".md", ".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        else:
            print(f"Error: Unsupported file format '{ext}'. Supported: .pdf, .docx, .md, .txt")
            sys.exit(1)
    except Exception as e:
        print(f"Error reading {ext} file: {e}")
        sys.exit(1)
    return text


def extract_company_name(job_description: str) -> str:
    """Extract company name from job description using common patterns."""
    # Common patterns for company names
    patterns = [
        r"(?:at|@|for|with)\s+([A-Z][A-Za-z0-9\s&]+?)(?:\s+is|\s+we|\s*[,\.])",
        r"^([A-Z][A-Za-z0-9\s&]+?)\s+(?:is looking|is seeking|is hiring)",
        r"Company:\s*([A-Za-z0-9\s&]+)",
        r"About\s+([A-Z][A-Za-z0-9\s&]+?)(?:\s*:|\s*\n)",
    ]

    for pattern in patterns:
        match = re.search(pattern, job_description, re.MULTILINE)
        if match:
            company = match.group(1).strip()
            # Clean up common suffixes
            company = re.sub(r'\s+(Inc|LLC|Ltd|Corp|Corporation)\.?$', '', company, flags=re.IGNORECASE)
            if len(company) > 2 and len(company) < 50:
                return company

    return None


def get_job_title(job_description: str) -> str:
    """Extract job title from job description."""
    patterns = [
        r"(?:Position|Title|Role|Job Title):\s*([^\n]+)",
        r"^([A-Z][A-Za-z\s\-/]+?)\s*(?:\n|$)",
        r"hiring\s+(?:a|an)\s+([A-Za-z\s\-/]+?)(?:\s+to|\s+who|\s*\.)",
    ]

    for pattern in patterns:
        match = re.search(pattern, job_description, re.MULTILINE | re.IGNORECASE)
        if match:
            title = match.group(1).strip()
            if len(title) > 3 and len(title) < 100:
                return title

    return "Position"


def sanitize_folder_name(name: str) -> str:
    """Convert string to safe folder name."""
    # Remove or replace invalid characters
    safe_name = re.sub(r'[<>:"/\\|?*]', '', name)
    safe_name = re.sub(r'\s+', '_', safe_name)
    return safe_name[:50]  # Limit length


def create_output_folder(company_name: str, config: dict) -> Path:
    """Create and return the output folder path for a company."""
    base_dir = Path(__file__).parent / config.get("output_base_dir", "applications")

    if not company_name:
        company_name = f"Application_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

    folder_name = sanitize_folder_name(company_name)
    output_dir = base_dir / folder_name

    # Handle existing folders by adding timestamp
    if output_dir.exists():
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_dir = base_dir / f"{folder_name}_{timestamp}"

    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def optimize_resume_with_claude(resume_text: str, job_description: str, client: anthropic.Anthropic) -> str:
    """Reject the retired standalone rewrite helper before using its client."""
    del resume_text, job_description, client
    raise NativeResumeTeamRequiredError()


def generate_cover_letter_with_claude(resume_text: str, job_description: str, company_name: str, job_title: str, config: dict, client: anthropic.Anthropic) -> str:
    """Use Claude to generate a tailored cover letter."""

    user_name = config.get("user_name", "")

    prompt = f"""Role: Expert Career Coach and Professional Writer specializing in compelling cover letters.

Task: Write a persuasive one-page cover letter for the following job application.

Input Data:
1. Applicant Resume:
{resume_text}

2. Target Job Description:
{job_description}

3. Company Name: {company_name or "the company"}
4. Job Title: {job_title}
5. Applicant Name: {user_name}

Instructions:
1. **Opening Hook**: Start with a compelling statement that shows genuine interest and immediately highlights the strongest qualification match.

2. **Value Proposition**: In the body, connect 3-4 specific experiences from the resume to key requirements in the JD. Use the STAR method briefly (Situation, Task, Action, Result).

3. **Company Research Connection**: Reference something specific about the company/role that shows research and genuine interest.

4. **Cultural Fit**: Demonstrate alignment with company values or mission if mentioned in JD.

5. **Strong Close**: End with confidence, a clear call to action, and enthusiasm for the opportunity.

Format Requirements:
- Maximum ONE page (350-400 words)
- Professional but personable tone
- 4-5 paragraphs maximum
- Do NOT include placeholder brackets like [Your Address] - write as if ready to send
- Start directly with the greeting (Dear Hiring Manager, or Dear [Company] Team,)

Output the complete cover letter text only, ready to be formatted."""

    message = client.messages.create(
        model="claude-sonnet-5",
        max_tokens=2000,
        thinking={"type": "disabled"},
        messages=[{"role": "user", "content": prompt}]
    )

    return message.content[0].text


def save_as_docx(
    content: str, output_path: Path, doc_type: str = "resume", *,
    master_path: str | Path | None = None, config_path: str = "config.json"
):
    """Save content as a formatted DOCX file."""
    destination = Path(output_path)
    temp_docx = None
    temp_md = None
    if doc_type == "resume":
        from resume_integrity_audit import (
            audit_resume_files,
            format_audit_report,
            resolve_master_path,
        )

        resolved_master = Path(master_path) if master_path is not None else resolve_master_path(config_path)
        destination.parent.mkdir(parents=True, exist_ok=True)
        md_handle = tempfile.NamedTemporaryFile(
            mode="w", encoding="utf-8", prefix=f".{destination.name}.",
            suffix=".tmp.md", dir=destination.parent, delete=False,
        )
        temp_md = Path(md_handle.name)
        try:
            md_handle.write(content)
        finally:
            md_handle.close()
        pre_report = audit_resume_files(
            tailored_path=temp_md, master_path=resolved_master, config_path=config_path
        )
        if not pre_report["passed"]:
            temp_md.unlink(missing_ok=True)
            raise ValueError(format_audit_report(pre_report))
        temp_md.unlink(missing_ok=True)
        docx_handle = tempfile.NamedTemporaryFile(
            prefix=f".{destination.name}.", suffix=".tmp.docx",
            dir=destination.parent, delete=False,
        )
        temp_docx = Path(docx_handle.name)
        docx_handle.close()
        output_path = temp_docx

    doc = Document()

    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Parse and add content
    lines = content.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Handle headers (markdown style or ALL CAPS)
        if line.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## ') or line.isupper():
            p = doc.add_paragraph()
            text = line[3:] if line.startswith('## ') else line
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.bold = True
            run.font.size = Pt(11)
        elif line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
            # Bullet points
            text = line[2:] if line.startswith(('- ', '* ')) else line[2:]
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
        elif line.startswith('**') and line.endswith('**'):
            # Bold text
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        else:
            # Regular paragraph
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)

    try:
        doc.save(output_path)
        if doc_type == "resume":
            post_report = audit_resume_files(
                tailored_path=temp_docx, master_path=resolved_master, config_path=config_path
            )
            if not post_report["passed"]:
                raise ValueError(format_audit_report(post_report))
            os.replace(temp_docx, destination)
            temp_docx = None
        print(f"Saved: {destination}")
    finally:
        if temp_md is not None:
            temp_md.unlink(missing_ok=True)
        if temp_docx is not None:
            temp_docx.unlink(missing_ok=True)


def get_multiline_input(prompt_text: str) -> str:
    """Get multi-line input from user."""
    print(prompt_text)
    print("(Enter your text, then type 'END' on a new line when done)")
    print("-" * 50)

    lines = []
    while True:
        try:
            line = input()
            if line.strip().upper() == 'END':
                break
            lines.append(line)
        except EOFError:
            break

    return '\n'.join(lines)


def main():
    """Fail closed before configuration, model, directory, or file access."""
    print("\n" + "=" * 60)
    print("       RESUME BUILDER - LEGACY DIRECT REWRITE DISABLED")
    print("=" * 60 + "\n")
    print(f"{NATIVE_RESUME_TEAM_REQUIRED}: {NATIVE_RESUME_TEAM_REQUIRED_MESSAGE}")
    return 2


if __name__ == "__main__":
    raise SystemExit(main())
