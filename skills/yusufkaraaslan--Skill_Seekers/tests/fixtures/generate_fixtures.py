"""Generate synthetic test fixtures for scraper integration tests.

Creates minimal but valid files for PDF, DOCX, and EPUB formats
using the same libraries the scrapers use (PyMuPDF, python-docx, ebooklib).

Usage:
    python tests/fixtures/generate_fixtures.py
"""

import zipfile
from pathlib import Path

OUT_DIR = Path(__file__).parent / "synthetic"


def generate_pdf():
    """Generate a minimal 3-page PDF using PyMuPDF (core dependency)."""
    try:
        import fitz
    except ImportError:
        print("SKIP: PyMuPDF (fitz) not installed")
        return False

    path = OUT_DIR / "document.pdf"
    doc = fitz.open()
    page1 = doc.new_page()
    page1.insert_text(fitz.Point(72, 72), "Test PDF Document")
    page1.insert_text(fitz.Point(72, 100), "This is page 1 of a test document.")
    page1.insert_text(fitz.Point(72, 128), "It contains sample content for testing.")
    page1.insert_text(fitz.Point(72, 156), "def hello():")
    page1.insert_text(fitz.Point(88, 172), 'print("Hello, World!")')

    page2 = doc.new_page()
    page2.insert_text(fitz.Point(72, 72), "Chapter 1: Getting Started")
    page2.insert_text(fitz.Point(72, 100), "This chapter covers installation and basic setup.")
    page2.insert_text(fitz.Point(72, 128), "## Installation")
    page2.insert_text(fitz.Point(72, 144), "pip install example-package")
    page2.insert_text(fitz.Point(72, 170), "import example_package")
    page2.insert_text(fitz.Point(72, 186), "result = example_package.do_thing()")

    page3 = doc.new_page()
    page3.insert_text(fitz.Point(72, 72), "Chapter 2: Advanced Usage")
    page3.insert_text(
        fitz.Point(72, 100), "This chapter covers advanced features and configuration."
    )
    page3.insert_text(fitz.Point(72, 128), "| Setting | Default | Description |")
    page3.insert_text(fitz.Point(72, 144), "| timeout | 30 | Request timeout in seconds |")
    page3.insert_text(fitz.Point(72, 160), "| retries | 3 | Number of retry attempts |")

    doc.save(str(path))
    doc.close()
    print(f"  Created: {path.name} ({path.stat().st_size} bytes)")
    return True


def generate_docx():
    """Generate a minimal DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        print("SKIP: python-docx not installed")
        return False

    path = OUT_DIR / "document.docx"
    doc = Document()
    doc.core_properties.title = "Test Word Document"
    doc.core_properties.subject = "Testing Skill Seekers DOCX scraper integration"

    doc.add_heading("Test Word Document", level=1)
    doc.add_paragraph("This is a test document for the Skill Seekers Word scraper.")

    doc.add_heading("Installation", level=2)
    doc.add_paragraph("pip install example-package")
    doc.add_paragraph("import example_package")

    doc.add_heading("Code Example", level=2)
    doc.add_paragraph(
        "def process_data(items):\n"
        "    results = []\n"
        "    for item in items:\n"
        "        results.append(item * 2)\n"
        "    return results"
    )

    table = doc.add_table(rows=3, cols=2, style="Light Grid Accent 1")
    table.cell(0, 0).text = "Setting"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "timeout"
    table.cell(1, 1).text = "30"
    table.cell(2, 0).text = "retries"
    table.cell(2, 1).text = "3"

    doc.save(str(path))
    print(f"  Created: {path.name} ({path.stat().st_size} bytes)")
    return True


def generate_epub():
    """Generate a minimal EPUB using ebooklib.

    Falls back to manual EPUB construction if ebooklib is not installed.
    """
    try:
        from ebooklib import epub
    except ImportError:
        return _generate_epub_manual()

    path = OUT_DIR / "document.epub"
    book = epub.EpubBook()
    book.set_identifier("test-epub-001")
    book.set_title("Test EPUB Document")
    book.set_language("en")
    book.add_author("Skill Seekers Test")
    book.add_metadata("DC", "description", "Testing Skill Seekers EPUB scraper integration")

    c1 = epub.EpubHtml(
        title="Chapter 1: Getting Started",
        file_name="chap1.xhtml",
        lang="en",
    )
    c1.content = (
        "<h1>Chapter 1: Getting Started</h1>"
        "<p>This chapter covers installation and basic setup.</p>"
        "<h2>Installation</h2>"
        "<pre><code>pip install example-package</code></pre>"
        "<p>Then import and use:</p>"
        "<pre><code>import example_package\n"
        "result = example_package.do_thing()</code></pre>"
    )

    c2 = epub.EpubHtml(
        title="Chapter 2: Advanced Usage",
        file_name="chap2.xhtml",
        lang="en",
    )
    c2.content = (
        "<h1>Chapter 2: Advanced Usage</h1>"
        "<p>This chapter covers advanced features.</p>"
        "<h2>Configuration</h2>"
        "<table><tr><th>Setting</th><th>Default</th></tr>"
        "<tr><td>timeout</td><td>30</td></tr>"
        "<tr><td>retries</td><td>3</td></tr></table>"
        "<h2>Example</h2>"
        "<pre><code>def process(x):\n"
        "    return x * 2</code></pre>"
    )

    book.add_item(c1)
    book.add_item(c2)
    book.toc = [
        epub.Link("chap1.xhtml", "Chapter 1", "chap1"),
        epub.Link("chap2.xhtml", "Chapter 2", "chap2"),
    ]
    book.spine = ["nav", c1, c2]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    epub.write_epub(str(path), book)
    print(f"  Created: {path.name} ({path.stat().st_size} bytes)")
    return True


def _generate_epub_manual():
    """Generate a minimal valid EPUB v3 manually (ZIP + mimetype + required structure)."""
    path = OUT_DIR / "document.epub"

    container_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">'
        "<rootfiles>"
        '<rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>'
        "</rootfiles>"
        "</container>"
    )

    content_opf = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="book-id" version="3.0">'
        '<metadata xmlns:dc="http://purl.org/dc/elements/1.1/">'
        "<dc:title>Test EPUB Document</dc:title>"
        '<dc:creator id="author">Skill Seekers Test</dc:creator>'
        "<dc:language>en</dc:language>"
        '<dc:identifier id="book-id">test-epub-001</dc:identifier>'
        "<dc:description>Testing Skill Seekers EPUB scraper integration</dc:description>"
        "</metadata>"
        "<manifest>"
        '<item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>'
        '<item id="ch1" href="chap1.xhtml" media-type="application/xhtml+xml"/>'
        '<item id="ch2" href="chap2.xhtml" media-type="application/xhtml+xml"/>'
        '<item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>'
        "</manifest>"
        '<spine toc="ncx">'
        '<itemref idref="ch1"/>'
        '<itemref idref="ch2"/>'
        "</spine>"
        "</package>"
    )

    chap1_xhtml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<!DOCTYPE html><html xmlns="http://www.w3.org/1999/xhtml"><head><title>Chapter 1</title></head>'
        "<body>"
        "<h1>Chapter 1: Getting Started</h1>"
        "<p>This chapter covers installation and basic setup.</p>"
        "<h2>Installation</h2>"
        "<pre><code>pip install example-package</code></pre>"
        "<p>Then import and use:</p>"
        "<pre><code>import example_package\nresult = example_package.do_thing()</code></pre>"
        "</body></html>"
    )

    chap2_xhtml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<!DOCTYPE html><html xmlns="http://www.w3.org/1999/xhtml"><head><title>Chapter 2</title></head>'
        "<body>"
        "<h1>Chapter 2: Advanced Usage</h1>"
        "<p>This chapter covers advanced features.</p>"
        "<h2>Configuration</h2>"
        "<table><tr><th>Setting</th><th>Default</th></tr>"
        "<tr><td>timeout</td><td>30</td></tr>"
        "<tr><td>retries</td><td>3</td></tr></table>"
        "<h2>Example</h2>"
        "<pre><code>def process(x):\n    return x * 2</code></pre>"
        "</body></html>"
    )

    nav_xhtml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<!DOCTYPE html><html xmlns="http://www.w3.org/1999/xhtml" '
        'xmlns:epub="http://www.idpf.org/2007/ops">'
        "<head><title>Table of Contents</title></head>"
        "<body>"
        '<nav epub:type="toc">'
        "<ol>"
        '<li><a href="chap1.xhtml">Chapter 1</a></li>'
        '<li><a href="chap2.xhtml">Chapter 2</a></li>'
        "</ol>"
        "</nav>"
        "</body></html>"
    )

    toc_ncx = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">'
        "<head>"
        '<meta name="dtb:uid" content="test-epub-001"/>'
        "</head>"
        "<docTitle><text>Test EPUB</text></docTitle>"
        "<navMap>"
        '<navPoint id="ch1" playOrder="1"><navLabel><text>Chapter 1</text></navLabel>'
        '<content src="chap1.xhtml"/></navPoint>'
        '<navPoint id="ch2" playOrder="2"><navLabel><text>Chapter 2</text></navLabel>'
        '<content src="chap2.xhtml"/></navPoint>'
        "</navMap>"
        "</ncx>"
    )

    with zipfile.ZipFile(str(path), "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)
        zf.writestr("META-INF/container.xml", container_xml)
        zf.writestr("OEBPS/content.opf", content_opf)
        zf.writestr("OEBPS/chap1.xhtml", chap1_xhtml)
        zf.writestr("OEBPS/chap2.xhtml", chap2_xhtml)
        zf.writestr("OEBPS/nav.xhtml", nav_xhtml)
        zf.writestr("OEBPS/toc.ncx", toc_ncx)

    print(f"  Created: {path.name} ({path.stat().st_size} bytes) [manual EPUB]")
    return True


def main():
    print("Generating synthetic test fixtures...")
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    results = {
        "pdf": generate_pdf(),
        "docx": generate_docx(),
        "epub": generate_epub(),
    }
    ok = sum(results.values())
    total = len(results)
    print(f"\nDone: {ok}/{total} fixture types generated.")


if __name__ == "__main__":
    main()
